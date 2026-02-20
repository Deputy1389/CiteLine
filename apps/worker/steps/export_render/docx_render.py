"""
DOCX rendering for chronology export.

Handles generation of professional DOCX output for paralegal use.
Extracted from step12_export.py during refactor - behavior preserved exactly.
"""
from __future__ import annotations

import io
from datetime import date, datetime, timezone
from typing import TYPE_CHECKING

from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from apps.worker.steps.export_render.common import (
    _clean_narrative_text,
    _date_str,
    _pages_ref,
    _provider_name,
    _set_cell_shading,
)

if TYPE_CHECKING:
    from packages.shared.models import Event, Gap, Provider


def generate_docx(
    run_id: str,
    matter_title: str,
    events: list[Event],
    gaps: list[Gap],
    providers: list[Provider],
    page_map: dict[int, tuple[str, int]] | None = None,
    narrative_synthesis: str | None = None,
) -> bytes:
    """
    Generate a professional DOCX chronology for paralegal use.
    """
    doc = DocxDocument()

    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    title_para = doc.add_heading("CiteLine Chronology", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_lines = [
        f"Matter: {matter_title}",
        f"Run ID: {run_id}",
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_paragraph()

    if narrative_synthesis:
        doc.add_heading("Medical Chronology Analysis", level=1)
        narrative_para = doc.add_paragraph(_clean_narrative_text(narrative_synthesis))
        narrative_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in narrative_para.runs:
            run.font.size = Pt(10)
        doc.add_paragraph()

    dated_events = []
    undated_events = []
    flagged_events = []

    for evt in events:
        has_date = (evt.date is not None and evt.date.value is not None)
        needs_review = any(f in evt.flags for f in ("MISSING_DATE", "MISSING_SOURCE", "NEEDS_REVIEW", "LOW_CONFIDENCE"))
        if needs_review: flagged_events.append(evt)
        elif has_date: dated_events.append(evt)
        else: undated_events.append(evt)

    dated_events.sort(key=lambda e: e.date.sort_key() if e.date else (date.min, 0))

    total = len(events)
    dated_count = len(dated_events)
    flagged_count = len(flagged_events)
    undated_count = len(undated_events)
    pct_dated = f"{(dated_count / total * 100):.0f}%" if total > 0 else "N/A"

    doc.add_heading("Summary", level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats = [
        ("Total Events", str(total)),
        ("Dated Events", f"{dated_count} ({pct_dated})"),
        ("Undated Events", str(undated_count)),
        ("Flagged (Needs Review)", str(flagged_count)),
    ]
    for i, (label, val) in enumerate(stats):
        summary_table.cell(i, 0).text = label
        summary_table.cell(i, 1).text = val
        for cell in (summary_table.cell(i, 0), summary_table.cell(i, 1)):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    def _add_events_table(heading: str, event_list: list[Event]):
        if not event_list: return
        doc.add_heading(heading, level=1)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True
        headers = ["Date", "Provider", "Type", "Description", "Citation"]
        hdr_row = tbl.rows[0]
        for idx, hdr_text in enumerate(headers):
            cell = hdr_row.cells[idx]
            cell.text = hdr_text
            _set_cell_shading(cell, "2C3E50")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for evt in event_list:
            row = tbl.add_row()
            cells = row.cells
            cells[0].text = _date_str(evt)
            cells[1].text = _provider_name(evt, providers)
            cells[2].text = evt.event_type.value.replace("_", " ").title()
            facts = "\n".join(f"• {f.text}" for f in evt.facts[:6])
            cells[3].text = facts
            cells[4].text = _pages_ref(evt, page_map)
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    _add_events_table("Chronology", dated_events)
    review_events = undated_events + flagged_events
    if review_events:
        _add_events_table("Undated / Needs Review", review_events)
        doc.add_paragraph()
        for evt in flagged_events:
            flags_str = ", ".join(evt.flags) if evt.flags else "UNDATED"
            p = doc.add_paragraph(f"⚠ {evt.event_id}: {flags_str}", style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    if gaps:
        doc.add_heading("Appendix: Treatment Gaps", level=1)
        for gap in gaps:
            doc.add_paragraph(f"• {gap.start_date} → {gap.end_date} ({gap.duration_days} days)", style="List Bullet")

    doc.add_paragraph()
    disclaimer = doc.add_paragraph("Factual extraction with citations. Requires human review. This document does not constitute legal or medical advice.")
    for run in disclaimer.runs:
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
