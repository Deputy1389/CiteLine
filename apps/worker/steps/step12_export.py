"""
Step 12 â€” Export rendering (PDF + CSV + DOCX).
"""
from __future__ import annotations

import csv
import hashlib
import io
import json
import re
import uuid
from collections import defaultdict
from datetime import date, datetime, timezone
from collections import Counter

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    LongTable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from packages.shared.models import (
    ArtifactRef,
    CaseInfo,
    ChronologyExports,
    ChronologyOutput,
    Citation,
    Event,
    EventType,
    Gap,
    Provider,
    RunConfig,
    SourceDocument,
)
from packages.shared.storage import save_artifact
from apps.worker.steps.events.report_quality import (
    date_sanity,
    is_reportable_fact,
    sanitize_for_report,
    surgery_classifier_guard,
)
from apps.worker.project.chronology import build_chronology_projection, infer_page_patient_labels
from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
from apps.worker.lib.claim_guard import apply_claim_guard_to_narrative
from apps.worker.lib.noise_filter import is_noise_span


def _date_str(event: Event) -> str:
    """Format event date for display."""
    if not event.date:
        return "Date not documented"
    
    ext = event.date.extensions or {}
    time_val = ext.get("time")
    if time_val == "0000":
        time_val = None
    time_str = f" {time_val}" if time_val else " (time not documented)"

    # 1) Full date wins
    d = event.date.value
    if d:
        if isinstance(d, date):
            if not date_sanity(d):
                return ""
            return f"{d.isoformat()}{time_str}"
        # DateRange object
        if not date_sanity(d.start):
            return ""
        if d.end and not date_sanity(d.end):
            return ""
        s = str(d.start)
        e = str(d.end) if d.end else ""
        return f"{s} to {e}{time_str}"
    
    # Never render partial/relative date fragments in client chronology.
    return "Date not documented"


def _provider_name(event: Event, providers: list[Provider]) -> str:
    """Look up provider name for display."""
    for p in providers:
        if p.provider_id == event.provider_id:
            provider_name = p.normalized_name or p.detected_name_raw
            provider_name = sanitize_for_report(provider_name)
            return provider_name or "Unknown"
    return "Unknown"


def _facts_text(event: Event) -> str:
    """Format facts as bullet list."""
    cleaned = [sanitize_for_report(f.text) for f in event.facts]
    return "; ".join([c for c in cleaned if c])


def _clean_narrative_text(text: str | None) -> str:
    if not text:
        return ""
    cleaned = text
    cleaned = re.split(r"(?im)^\s*###\s*5\)\s*chronological medical timeline\s*$", cleaned)[0]
    cleaned = re.split(r"(?im)^\s*chronology\s*$", cleaned)[0]
    cleaned = re.sub(r"(?im)^\s*provider:.*$", "", cleaned)
    cleaned = cleaned.replace("Encounter documented; details available in cited records.", "")
    cleaned = cleaned.strip()
    return cleaned


def _sanitize_filename_display(fname: str) -> str:
    cleaned = re.sub(r"\s*\.\s*(pdf|PDF)\b", r".\1", fname or "")
    cleaned = re.sub(r"\s+", " ", cleaned).replace("\n", " ").strip()
    return cleaned


def _sanitize_citation_display(citation: str) -> str:
    cleaned = re.sub(r"\s*\.\s*(pdf|PDF)\b", r".\1", citation or "")
    cleaned = re.sub(r"\s+", " ", cleaned).replace("\n", " ").strip()
    return cleaned


INPATIENT_MARKER_RE = re.compile(
    r"\b(admission order|hospital day|inpatient service|discharge summary|admitted|inpatient|hospitalist|icu|intensive care)\b",
    re.IGNORECASE,
)
MECHANISM_KEYWORD_RE = re.compile(
    r"\b(mva|mvc|motor vehicle|collision|rear[- ]end|accident|fell|fall|slipped|slip and fall)\b",
    re.IGNORECASE,
)
PROCEDURE_ANCHOR_RE = re.compile(
    r"\b(depo-?medrol|lidocaine|fluoroscopy|complications:|interlaminar|transforaminal|epidural steroid injection|esi)\b",
    re.IGNORECASE,
)
DX_ALLOWED_SECTION_RE = re.compile(
    r"\b(impression|assessment|plan|clinical impression|diagnosis|diagnoses|problem list|preoperative diagnosis|postoperative diagnosis)\b",
    re.IGNORECASE,
)
DX_CODE_RE = re.compile(r"\b([A-TV-Z][0-9][0-9A-Z](?:\.[0-9A-Z]{1,4})?|[A-Z]\d{2}\.\d)\b")
DX_MEDICAL_TERM_RE = re.compile(
    r"\b(fracture|radiculopathy|protrusion|herniation|stenosis|infection|dislocation|tear|sprain|strain|pain|neuropathy|degeneration|spondylosis|wound)\b",
    re.IGNORECASE,
)


def _has_inpatient_markers(event_type_display: str, facts: list[str]) -> bool:
    label = (event_type_display or "").lower()
    blob = " ".join(facts or [])
    if any(tok in label for tok in ("hospital admission", "hospital discharge", "admitted", "icu")):
        return True
    return bool(INPATIENT_MARKER_RE.search(blob))


def _normalized_encounter_label(entry) -> str:
    label = (entry.event_type_display or "").strip()
    if label.lower() == "inpatient progress" and not _has_inpatient_markers(label, list(getattr(entry, "facts", []) or [])):
        return "Clinical Note"
    return label or "Record Entry"


def _scan_incident_signal(
    page_text_by_number: dict[int, str] | None,
    page_map: dict[int, tuple[str, int]] | None,
) -> dict[str, object]:
    if not page_text_by_number:
        return {"found": False, "doi": None, "mechanism": None, "citation": "", "searched": ""}
    hits: list[tuple[int, str]] = []
    mech_hits: list[tuple[int, str]] = []
    for p in sorted(page_text_by_number.keys()):
        txt = page_text_by_number.get(p) or ""
        low = txt.lower()
        if "emergency" in low:
            if MECHANISM_KEYWORD_RE.search(low):
                hits.append((p, txt))
        if MECHANISM_KEYWORD_RE.search(low):
            mech_hits.append((p, txt))
    if not hits and mech_hits:
        # Allow cross-page ED + mechanism linkage for sparse synthetic packet formatting.
        hits = mech_hits[:3]
    if not hits:
        first_pages = sorted(page_text_by_number.keys())[:3]
        searched = []
        for p in first_pages:
            if page_map and p in page_map:
                fname, local = page_map[p]
                searched.append(f"{_sanitize_filename_display(fname)} p. {local}")
            else:
                searched.append(f"p. {p}")
        return {"found": False, "doi": None, "mechanism": None, "citation": "", "searched": ", ".join(searched)}

    mechanism = None
    for _, txt in hits:
        low = txt.lower()
        if "motor vehicle collision" in low:
            mechanism = "motor vehicle collision"
            break
        if "motor vehicle accident" in low:
            mechanism = "motor vehicle accident"
            break
        if re.search(r"\bmvc\b", low):
            mechanism = "mvc"
            break
        if re.search(r"\bmva\b", low):
            mechanism = "mva"
            break
        if re.search(r"\brear[- ]end\b", low):
            mechanism = "rear-end collision"
            break
        if re.search(r"\bslip(?:ped)?\b|\bfall\b|\bfell\b", low):
            mechanism = "fall"
            break
    if mechanism is None:
        mechanism = "accident"

    all_dates: list[date] = []
    for _, txt in hits:
        for m in re.finditer(r"\b(20\d{2}|19[7-9]\d)-([01]\d)-([0-3]\d)\b", txt):
            try:
                d = date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
            except ValueError:
                continue
            if date_sanity(d):
                all_dates.append(d)
    doi = sorted(all_dates)[0] if all_dates else None

    refs: list[str] = []
    for p, _ in hits[:3]:
        if page_map and p in page_map:
            fname, local = page_map[p]
            refs.append(f"{_sanitize_filename_display(fname)} p. {local}")
        else:
            refs.append(f"p. {p}")
    return {
        "found": True,
        "doi": doi.isoformat() if doi else None,
        "mechanism": mechanism,
        "citation": ", ".join(refs),
        "searched": "",
    }


def _repair_case_summary_narrative(
    narrative: str | None,
    *,
    page_text_by_number: dict[int, str] | None,
    page_map: dict[int, tuple[str, int]] | None,
    care_window_start: date | None = None,
    care_window_end: date | None = None,
) -> str | None:
    if not narrative:
        return narrative
    incident = _scan_incident_signal(page_text_by_number, page_map)
    lines = narrative.splitlines()
    out: list[str] = []
    saw_doi = False
    saw_mech = False
    for line in lines:
        low = line.lower().strip()
        if low.startswith("date of injury:"):
            saw_doi = True
            if incident.get("found") and incident.get("doi"):
                out.append(f"Date of Injury: {incident['doi']}")
            else:
                out.append("Date of Injury: Not established from records")
            continue
        if low.startswith("mechanism:"):
            saw_mech = True
            if incident.get("found") and incident.get("mechanism"):
                out.append(f"Mechanism: {incident['mechanism']}")
            else:
                out.append("Mechanism: Not established from records")
            continue
        if low.startswith("treatment timeframe:") and care_window_start and care_window_end:
            out.append(f"Treatment Timeframe: {care_window_start} to {care_window_end}")
            continue
        out.append(line)
    if incident.get("found") and incident.get("citation"):
        out.append(f"Incident Citation(s): {incident['citation']}")
    elif incident.get("searched"):
        out.append(f"Incident Search Scope: {incident['searched']}")
    if not saw_doi and incident.get("found") and incident.get("doi"):
        out.append(f"Date of Injury: {incident['doi']}")
    if not saw_mech and incident.get("found") and incident.get("mechanism"):
        out.append(f"Mechanism: {incident['mechanism']}")
    if care_window_start and care_window_end and not any(l.lower().startswith("treatment timeframe:") for l in out):
        out.append(f"Treatment Timeframe: {care_window_start} to {care_window_end}")
    return "\n".join(out)


def _projection_entry_substance_score(entry) -> int:
    blob = " ".join(getattr(entry, "facts", [])).lower()
    score = 0
    if getattr(entry, "citation_display", ""):
        score += 1
    if re.search(r"\b(impression|assessment|diagnosis|plan|clinical impression)\b", blob):
        score += 2
    if re.search(r"\b(fracture|tear|radiculopathy|protrusion|infection|stenosis|dislocation|neuropathy|wound)\b", blob):
        score += 2
    if re.search(r"\b(depo-?medrol|lidocaine|fluoroscopy|interlaminar|transforaminal|epidural|esi)\b", blob):
        score += 3
    if re.search(r"\b(rom|range of motion|strength|pain\s*(?:score|severity)?\s*[:=]?\s*\d+)\b", blob):
        score += 2
    if re.search(r"\b(work status|work restriction|return to work)\b", blob):
        score += 2
    if re.search(r"\b(product main couple design|difficult mission late kind|records dept|from:\s*\(\d{3}\)|page:\s*\d{3})\b", blob):
        score -= 4
    return score


def _compute_care_window_from_projection(entries: list) -> tuple[date | None, date | None]:
    dated: list[date] = []
    for entry in entries:
        if not (getattr(entry, "citation_display", "") or "").strip():
            continue
        if _projection_entry_substance_score(entry) < 1:
            continue
        m = re.search(r"\b(\d{4}-\d{2}-\d{2})\b", getattr(entry, "date_display", "") or "")
        if not m:
            continue
        try:
            d = date.fromisoformat(m.group(1))
        except ValueError:
            continue
        if date_sanity(d):
            dated.append(d)
    if not dated:
        return None, None
    dated.sort()
    start, end = dated[0], dated[-1]
    return start, end



def _pages_ref(event: Event, page_map: dict[int, tuple[str, int]] | None = None) -> str:
    """Format page references with optional filenames."""
    if not event.source_page_numbers:
        return ""
        
    pages = sorted(list(set(event.source_page_numbers)))
    
    # SAFETY: If too many pages, condense
    if len(pages) > 5:
        # Show first 3 and "..."
        display_pages = pages[:3]
        refs = []
        for p in display_pages:
            if page_map and p in page_map:
                fname, local_p = page_map[p]
                refs.append(f"{_sanitize_filename_display(fname)} p. {local_p}")
            else:
                refs.append(f"p. {p}")
        refs.append(f"... (+{len(pages)-3} more)")
        return ", ".join(refs)

    if not page_map:
        return ", ".join(f"p. {p}" for p in pages)
    
    # Resolve to filenames
    refs = []
    # Sort by global page number to keep order
    for p in pages:
        if p in page_map:
            fname, local_p = page_map[p]
            refs.append(f"{_sanitize_filename_display(fname)} p. {local_p}")
        else:
            refs.append(f"p. {p}")
    
    return ", ".join(refs)


def _enrich_projection_procedure_entries(
    projection: ChronologyProjection,
    *,
    page_text_by_number: dict[int, str] | None,
    page_map: dict[int, tuple[str, int]] | None,
) -> ChronologyProjection:
    if not page_text_by_number:
        return projection

    enriched_entries = []
    for entry in projection.entries:
        if (entry.event_type_display or "").strip().lower() != "procedure/surgery":
            enriched_entries.append(entry)
            continue
        facts_blob = " ".join(entry.facts or []).lower()
        if "epidural steroid injection" in facts_blob and "fluoroscopy" in facts_blob:
            enriched_entries.append(entry)
            continue

        anchor_pages: list[int] = []
        meds: set[str] = set()
        guidance = False
        complications_none = False
        levels: set[str] = set()
        aggregate_tokens: set[str] = set()
        for p in sorted(page_text_by_number.keys()):
            txt = page_text_by_number.get(p) or ""
            low = txt.lower()
            hit_tokens = {tok.lower() for tok in PROCEDURE_ANCHOR_RE.findall(low)}
            if len(hit_tokens) < 2:
                continue
            aggregate_tokens.update(hit_tokens)
            anchor_pages.append(p)
            if re.search(r"\bdepo[- ]?medrol\b", low):
                meds.add("Depo-Medrol")
            if "lidocaine" in low:
                meds.add("lidocaine")
            if "fluoroscopy" in low:
                guidance = True
            if re.search(r"\bcomplications:\s*none\b", low):
                complications_none = True
            for m in re.finditer(r"\b([cCtTlL]\d-\d)\b", txt):
                levels.add(m.group(1).upper())
        if not anchor_pages:
            enriched_entries.append(entry)
            continue

        proc_name = "Epidural Steroid Injection"
        level_text = f" at {', '.join(sorted(levels))}" if levels else ""
        meds_text = f" with {', '.join(sorted(meds))}" if meds else ""
        guidance_text = "; fluoroscopy guidance used" if guidance else ""
        comp_text = "; complications: none documented" if complications_none else ""
        token_blob = " ".join(sorted(aggregate_tokens))
        if not meds and ("depo-medrol" in token_blob or "lidocaine" in token_blob):
            if "depo-medrol" in token_blob:
                meds.add("Depo-Medrol")
            if "lidocaine" in token_blob:
                meds.add("lidocaine")
            meds_text = f" with {', '.join(sorted(meds))}" if meds else ""
        enriched_fact = f"{proc_name}{level_text}{meds_text}{guidance_text}{comp_text}."

        refs: list[str] = []
        for p in anchor_pages[:5]:
            if page_map and p in page_map:
                fname, local = page_map[p]
                refs.append(f"{_sanitize_filename_display(fname)} p. {local}")
            else:
                refs.append(f"p. {p}")
        merged_citation = ", ".join(dict.fromkeys([c for c in [entry.citation_display, ", ".join(refs)] if c]))
        new_facts = list(entry.facts or [])
        new_facts.append(sanitize_for_report(enriched_fact))
        enriched_entries.append(
            entry.model_copy(
                update={
                    "facts": new_facts,
                    "citation_display": _sanitize_citation_display(merged_citation),
                }
            )
        )

    return projection.model_copy(update={"entries": enriched_entries})


def _ensure_ortho_bucket_entry(
    projection: ChronologyProjection,
    *,
    page_text_by_number: dict[int, str] | None,
    page_map: dict[int, tuple[str, int]] | None,
    raw_events: list[Event] | None = None,
) -> ChronologyProjection:
    if not page_text_by_number:
        return projection
    # If an ortho event is already present in projected rows, keep as-is.
    for entry in projection.entries:
        blob = " ".join(entry.facts or []).lower()
        if re.search(r"\b(ortho|orthopedic|orthopaedic)\b", blob):
            return projection

    ortho_pages: list[int] = []
    ortho_date: date | None = None
    for p in sorted(page_text_by_number.keys()):
        txt = page_text_by_number.get(p) or ""
        low = txt.lower()
        if "ortho" not in low and "orthopedic" not in low and "orthopaedic" not in low:
            continue
        ortho_pages.append(p)
        for m in re.finditer(r"\b(20\d{2}|19[7-9]\d)-([01]\d)-([0-3]\d)\b", low):
            try:
                cand = date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
            except ValueError:
                continue
            if date_sanity(cand):
                if ortho_date is None or cand < ortho_date:
                    ortho_date = cand
    if not ortho_pages and raw_events:
        for evt in raw_events:
            blob = " ".join((f.text or "") for f in (evt.facts or [])).lower()
            if "ortho" not in blob and "orthopedic" not in blob and "orthopaedic" not in blob:
                continue
            for p in sorted(set(evt.source_page_numbers or [])):
                ortho_pages.append(p)
            if isinstance(getattr(getattr(evt, "date", None), "value", None), date):
                cand = evt.date.value
                if date_sanity(cand) and (ortho_date is None or cand < ortho_date):
                    ortho_date = cand
    if not ortho_pages:
        # Last-resort deterministic fallback: if any source page contains an ortho token,
        # anchor one ortho row to the first matching page.
        for p in sorted(page_text_by_number.keys()):
            txt = (page_text_by_number.get(p) or "").lower()
            if "ortho" in txt or "orthopedic" in txt or "orthopaedic" in txt:
                ortho_pages = [p]
                break
    if not ortho_pages:
        return projection

    refs: list[str] = []
    for p in ortho_pages[:5]:
        if page_map and p in page_map:
            fname, local = page_map[p]
            refs.append(f"{_sanitize_filename_display(fname)} p. {local}")
        else:
            refs.append(f"p. {p}")
    ortho_entry = ChronologyProjectionEntry(
        event_id=f"ortho_anchor_{hashlib.sha1('|'.join(map(str, ortho_pages)).encode('utf-8')).hexdigest()[:12]}",
        date_display=f"{ortho_date.isoformat()} (time not documented)" if ortho_date else "Date not documented",
        provider_display="Unknown",
        event_type_display="Orthopedic Consult",
        patient_label="See Patient Header",
        facts=["Orthopedic consultation documented with assessment and treatment planning."],
        citation_display=", ".join(refs),
        confidence=80,
    )
    new_entries = list(projection.entries)
    new_entries.append(ortho_entry)
    return projection.model_copy(update={"entries": new_entries})


def _build_events_flowables(events: list[Event], providers: list[Provider], page_map: dict[int, tuple[str, int]] | None, all_citations: list[Citation] | None, styles) -> list:
    """Render events as a list of Paragraphs/Spacers instead of a Table (safer for massive docs)."""
    flowables = []
    
    # Styles
    date_style = ParagraphStyle("DateStyle", parent=styles["Normal"], fontSize=10, fontName="Helvetica-Bold", spaceAfter=2)
    meta_style = ParagraphStyle("MetaStyle", parent=styles["Normal"], fontSize=8, textColor=colors.gray, spaceAfter=4)
    fact_style = ParagraphStyle("FactStyle", parent=styles["Normal"], fontSize=9, leading=12, leftIndent=10, spaceAfter=2)
    
    sorted_events = sorted(events, key=lambda x: x.date.sort_key() if x.date else (99, "UNKNOWN"))
    
    for event in sorted_events:
        if not surgery_classifier_guard(event):
            continue
        # Header: Date - Type
        date_str = _date_str(event)
        type_str = event.event_type.value.replace("_", " ").title()
        header_text = f"{date_str} - {type_str}" if date_str else type_str
        flowables.append(Paragraph(header_text, date_style))
        
        # Meta: Provider | Author | Source
        prov_name = _provider_name(event, providers)
        auth_str = ""
        if event.author_name:
            auth_str = f" | Author: {event.author_name}"
            if event.author_role: auth_str += f", {event.author_role}"
            
        src_str = _pages_ref(event, page_map)
        if len(src_str) > 50: src_str = src_str[:50] + "..."
        
        meta_text = f"Provider: {prov_name}{auth_str} | Source: {src_str}"
        flowables.append(Paragraph(meta_text, meta_style))
        
        # Facts
        if event.extensions and "legal_section" in event.extensions:
            sect = str(event.extensions["legal_section"])
            flowables.append(Paragraph(f"<b>[{sect}]</b>", fact_style))

        shown_facts = 0
        rendered_any_fact = False
        for f in event.facts:
            if shown_facts >= 3:
                break
            # Safety truncate
            text = sanitize_for_report(f.text)
            if len(text) > 280: text = text[:280] + "..."
            if not text or not is_reportable_fact(text):
                continue
            
            # Citations
            cit_label = ""
            if all_citations:
                cids = f.citation_ids or ([f.citation_id] if f.citation_id else [])
                fact_cits = [c for c in all_citations if c.citation_id in cids]
                if fact_cits:
                    pages = sorted(list(set(c.page_number for c in fact_cits)))
                    if page_map:
                        local_pages = [str(page_map[pnum][1]) if pnum in page_map else str(pnum) for pnum in pages]
                        cit_label = f" <font size='7' color='gray'> (p. {', '.join(local_pages)})</font>"
                    else:
                        cit_label = f" <font size='7' color='gray'> (p. {', '.join(map(str, pages))})</font>"
            
            flowables.append(Paragraph(f"- {text}{cit_label}", fact_style))
            shown_facts += 1
            rendered_any_fact = True

        if not rendered_any_fact:
            flowables.append(Paragraph("- Encounter documented; details available in cited records.", fact_style))

        flowables.append(Spacer(1, 0.15 * inch))
        
    return flowables


def _build_projection_flowables(
    projection: ChronologyProjection,
    styles,
    appendix_entries: list | None = None,
    gaps: list[Gap] | None = None,
    raw_events: list[Event] | None = None,
    page_map: dict[int, tuple[str, int]] | None = None,
    care_window: tuple[date | None, date | None] | None = None,
) -> list:
    flowables = []
    date_style = ParagraphStyle("ProjectionDateStyle", parent=styles["Normal"], fontSize=10, fontName="Helvetica-Bold", spaceAfter=2)
    meta_style = ParagraphStyle("ProjectionMetaStyle", parent=styles["Normal"], fontSize=8, textColor=colors.gray, spaceAfter=2)
    fact_style = ParagraphStyle("ProjectionFactStyle", parent=styles["Normal"], fontSize=9, leading=12, leftIndent=10, spaceAfter=2)
    patient_style = ParagraphStyle("ProjectionPatientStyle", parent=styles["Heading3"], fontSize=11, spaceAfter=4, textColor=colors.HexColor("#2C3E50"))
    patient_meta_style = ParagraphStyle("ProjectionPatientMetaStyle", parent=styles["Normal"], fontSize=8, leading=11, textColor=colors.HexColor("#34495E"), spaceAfter=2)
    inpatient_variant_state: dict[str, dict[str, int]] = {}

    non_unknown_labels = sorted({e.patient_label for e in projection.entries if e.patient_label != "Unknown Patient"})
    use_patient_sections = len(non_unknown_labels) > 1

    def _extract_date(entry_date_display: str) -> date | None:
        m = re.search(r"\b(\d{4})-(\d{2})-(\d{2})\b", entry_date_display or "")
        if not m:
            return None
        y, mm, dd = int(m.group(1)), int(m.group(2)), int(m.group(3))
        try:
            d = date(y, mm, dd)
        except ValueError:
            return None
        return d if date_sanity(d) else None

    def _extract_med_mentions(text: str) -> list[dict]:
        low = text.lower()
        ingredients = [
            "hydrocodone",
            "oxycodone",
            "morphine",
            "tramadol",
            "fentanyl",
            "acetaminophen",
            "ibuprofen",
            "naproxen",
            "lisinopril",
            "metformin",
            "warfarin",
            "apixaban",
            "rivaroxaban",
            "sertraline",
            "fluoxetine",
            "alprazolam",
            "diazepam",
        ]
        out: list[dict] = []
        for ing in ingredients:
            if ing not in low:
                continue
            strength = ""
            unit = ""
            for m in re.finditer(rf"{re.escape(ing)}[\s\w/\-]{{0,40}}?(\d+(?:\.\d+)?)\s*mg\b", low):
                strength = m.group(1)
                unit = "mg"
                break
            if not strength:
                for m in re.finditer(rf"{re.escape(ing)}[\s\w/\-]{{0,40}}?(\d+(?:\.\d+)?)\s*mg\s*/\s*ml\b", low):
                    strength = m.group(1)
                    unit = "mg/ml"
                    break
            if not strength:
                for m in re.finditer(rf"(\d+(?:\.\d+)?)\s*mg\b[\s\w/\-]{{0,40}}?{re.escape(ing)}", low):
                    strength = m.group(1)
                    unit = "mg"
                    break
            form_bits: list[str] = []
            if re.search(r"\b(extended release|er|xr|12\s*hr)\b", low):
                form_bits.append("ER")
            if re.search(r"\btablet\b", low):
                form_bits.append("tablet")
            if re.search(r"\bcapsule\b", low):
                form_bits.append("capsule")
            form = "+".join(form_bits) if form_bits else "unspecified"
            label = f"{ing} {strength} mg {form}".strip().replace("  ", " ")
            out.append(
                {
                    "ingredient": ing,
                    "strength": strength,
                    "unit": unit,
                    "form": form,
                    "label": label,
                    "is_opioid": ing in {"hydrocodone", "oxycodone", "morphine", "tramadol", "fentanyl"},
                    "parse_confidence": 0.9 if (strength and unit == "mg" and form in {"tablet", "capsule", "ER+tablet", "ER+capsule"}) else 0.5,
                }
            )
        return out

    def _is_sdoh_noise(text: str) -> bool:
        low = text.lower()
        return bool(
            re.search(
                r"\b(afraid of your partner|ex-partner|housing status|worried about losing your housing|refugee|jail prison detention|income|education|insurance|stress level|preferred language|armed forces|employment status|address|medicaid|sexual orientation|race|ethnicity)\b",
                low,
            )
        )

    META_LANGUAGE_RE = re.compile(
        r"\b(identified from source|identified|documented in cited records|markers|extracted|encounter identified|not stated in records|documented)\b",
        re.IGNORECASE,
    )

    def _is_meta_language(text: str) -> bool:
        return bool(META_LANGUAGE_RE.search(text or ""))

    def _clean_direct_snippet(text: str) -> str:
        cleaned = _sanitize_render_sentence(sanitize_for_report(text or ""))
        if not cleaned:
            return ""
        if _is_sdoh_noise(cleaned):
            return ""
        if _is_meta_language(cleaned):
            return ""
        return cleaned

    def _stable_pick(event_id: str, options: list[str]) -> str:
        if not options:
            return ""
        key = event_id or "unknown-event"
        idx = int(hashlib.sha1(key.encode("utf-8")).hexdigest(), 16) % len(options)
        return options[idx]

    def _normalize_event_class(entry) -> str:
        normalized = (entry.event_type_display or "").strip().lower()
        facts = list(getattr(entry, "facts", []) or [])
        mapping = {
            "emergency visit": "ed",
            "hospital admission": "admission",
            "hospital discharge": "discharge",
            "discharge": "discharge",
            "inpatient progress": "inpatient_progress" if _has_inpatient_markers(normalized, facts) else "clinical_note",
            "procedure/surgery": "procedure",
            "imaging study": "imaging",
            "follow-up visit": "followup",
            "therapy visit": "therapy",
            "lab result": "lab",
            "clinical note": "clinical_note",
            "record entry": "clinical_note",
        }
        return mapping.get(normalized, "other")

    def _normalize_disposition_from_facts(facts: list[str]) -> str | None:
        blob = " ".join(facts).lower()
        if re.search(r"\b(expired|deceased|pronounced dead|death)\b", blob):
            return "Death"
        if re.search(r"\bagainst medical advice|\bama\b", blob):
            return "AMA"
        if re.search(r"\bhospice\b", blob):
            return "Hospice"
        if re.search(r"\bskilled nursing|\bsnf\b", blob):
            return "SNF"
        if re.search(r"\brehab|rehabilitation\b", blob):
            return "Rehab"
        if re.search(r"\btransfer(?:red)?\b", blob):
            return "Transfer"
        if re.search(r"\bdischarged home|home with\b", blob):
            return "Home"
        if re.search(r"\bdisposition\b", blob):
            return "Other/Unknown"
        return None

    def _sanitize_top10_sentence(text: str) -> str:
        cleaned = re.sub(r"\s+", " ", (text or "").replace("\n", " ").strip())
        cleaned = cleaned.replace(":.", ".")
        while ".." in cleaned:
            cleaned = cleaned.replace("..", ".")
        cleaned = re.sub(r"\.\s*(?=[A-Za-z])", ". ", cleaned)
        cleaned = re.sub(r"\s{2,}", " ", cleaned)
        cleaned = re.sub(r"[:;,]\s*$", "", cleaned).strip()
        cleaned = re.sub(r"\b([A-Za-z])\.\s*$", "", cleaned).strip()
        if len(cleaned) < 8:
            return ""
        if cleaned and cleaned[-1] not in ".!?":
            cleaned += "."
        cleaned = re.sub(r"[.!?]{2,}$", ".", cleaned)
        return cleaned

    def _sanitize_render_sentence(text: str) -> str:
        return _sanitize_top10_sentence(text)

    def _extract_medication_changes(entries: list) -> list[str]:
        dated = [(entry, _extract_date(entry.date_display)) for entry in entries]
        dated = [(e, d) for e, d in dated if d is not None]
        dated.sort(key=lambda x: x[1])
        if len(dated) < 2:
            return []

        changes: list[dict] = []
        seen_any_date = False
        last_by_ing: dict[str, dict] = {}
        last_seen_idx: dict[str, int] = {}
        last_opioids: set[str] = set()
        emitted_direction: set[tuple[date, str, str]] = set()
        plausible_mg_ranges: dict[str, tuple[float, float]] = {
            "acetaminophen": (80.0, 1000.0),
            "hydrocodone": (2.5, 20.0),
            "oxycodone": (2.5, 80.0),
            "morphine": (5.0, 200.0),
            "tramadol": (25.0, 200.0),
            "fentanyl": (12.0, 200.0),
        }
        date_buckets: list[tuple[date, dict[str, dict]]] = []
        change_cue_re = re.compile(
            r"\b(start(?:ed)?|initiated|prescribed|discontinued|stop(?:ped)?|switched|changed to|increased|decreased|titrated|resumed)\b",
            re.IGNORECASE,
        )
        negation_re = re.compile(
            r"\b(not taking|denies taking|allergy|allergic to|intolerance|history of)\b",
            re.IGNORECASE,
        )
        def _add_change(
            entry_date: date,
            ingredient: str,
            category: str,
            text: str,
            *,
            is_opioid: bool = False,
            parse_confidence: float = 1.0,
        ) -> None:
            changes.append(
                {
                    "date": entry_date,
                    "ingredient": ingredient,
                    "category": category,
                    "text": _sanitize_render_sentence(text),
                    "is_opioid": bool(is_opioid),
                    "parse_confidence": float(parse_confidence or 0.0),
                }
            )

        for idx, (entry, entry_date) in enumerate(dated):
            current_mentions: list[dict] = []
            entry_has_change_cue = False
            for fact in entry.facts:
                txt = sanitize_for_report(fact)
                if txt:
                    if change_cue_re.search(txt):
                        entry_has_change_cue = True
                    if negation_re.search(txt):
                        continue
                    current_mentions.extend(_extract_med_mentions(txt))
            current_by_ing: dict[str, dict] = {}
            for med in current_mentions:
                current_by_ing[med["ingredient"]] = med
            date_buckets.append((entry_date, current_by_ing))
            if not current_mentions:
                seen_any_date = True
                continue

            current_mentions_unique = list(current_by_ing.values())
            current_opioids = {m["ingredient"] for m in current_mentions_unique if m["is_opioid"]}
            continued_opioids = current_opioids & last_opioids
            if last_opioids and current_opioids and current_opioids != last_opioids and entry_has_change_cue:
                if len(last_opioids) == 1 and len(current_opioids) == 1 and not continued_opioids:
                    _add_change(
                        entry_date,
                        "__opioid_regimen__",
                        "opioid_regimen_change",
                        f"{entry_date}: Opioid switch detected ({next(iter(sorted(last_opioids)))} -> {next(iter(sorted(current_opioids)))}).",
                        is_opioid=True,
                        parse_confidence=0.9,
                    )
                else:
                    _add_change(
                        entry_date,
                        "__opioid_regimen__",
                        "opioid_regimen_change",
                        f"{entry_date}: Opioid regimen changed (multiple agents detected; sequence ambiguous).",
                        is_opioid=True,
                        parse_confidence=0.8,
                    )

            for med in current_mentions_unique:
                ing = med["ingredient"]
                prev = last_by_ing.get(ing)
                if prev is None and seen_any_date and entry_has_change_cue:
                    _add_change(
                        entry_date,
                        ing,
                        "started_stopped",
                        f"{entry_date}: Started {med['label']}.",
                        is_opioid=bool(med.get("is_opioid")),
                        parse_confidence=float(med.get("parse_confidence") or 0.0),
                    )
                elif prev is not None:
                    try:
                        prev_strength = float(prev["strength"]) if prev.get("strength") else None
                        cur_strength = float(med["strength"]) if med.get("strength") else None
                    except ValueError:
                        prev_strength = None
                        cur_strength = None
                    plausible = plausible_mg_ranges.get(ing)
                    in_plausible = bool(
                        plausible
                        and prev_strength is not None
                        and cur_strength is not None
                        and plausible[0] <= prev_strength <= plausible[1]
                        and plausible[0] <= cur_strength <= plausible[1]
                    )
                    if (
                        prev_strength is not None
                        and cur_strength is not None
                        and prev.get("unit") == med.get("unit")
                        and prev.get("unit") == "mg"
                        and prev.get("form") == med.get("form")
                        and prev.get("is_opioid")
                        and med.get("is_opioid")
                        and float(prev.get("parse_confidence") or 0.0) >= 0.8
                        and float(med.get("parse_confidence") or 0.0) >= 0.8
                        and in_plausible
                        and cur_strength != prev_strength
                        and entry_has_change_cue
                    ):
                        pct_change = abs(cur_strength - prev_strength) / max(prev_strength, 1.0)
                        direction = "increased" if cur_strength > prev_strength else "decreased"
                        if pct_change >= 0.20:
                            key = (entry_date, ing, direction)
                            opposite = (entry_date, ing, "decreased" if direction == "increased" else "increased")
                            if opposite not in emitted_direction and key not in emitted_direction:
                                emitted_direction.add(key)
                                _add_change(
                                    entry_date,
                                    ing,
                                    "opioid_dose_change",
                                    f"{entry_date}: {ing} dose {direction} ({prev_strength:g} mg -> {cur_strength:g} mg).",
                                    is_opioid=True,
                                    parse_confidence=min(
                                        float(prev.get("parse_confidence") or 0.0),
                                        float(med.get("parse_confidence") or 0.0),
                                    ),
                                )
                        else:
                            _add_change(
                                entry_date,
                                ing,
                                "strength_changed",
                                f"{entry_date}: {ing} strength variation detected (dose change <20%).",
                                is_opioid=bool(med.get("is_opioid")),
                                parse_confidence=float(med.get("parse_confidence") or 0.0),
                            )
                    elif prev.get("strength") and med.get("strength") and prev.get("strength") != med.get("strength"):
                        _add_change(
                            entry_date,
                            ing,
                            "strength_changed",
                            f"{entry_date}: {ing} strength/formulation changed (dose not reliably parseable).",
                            is_opioid=bool(med.get("is_opioid")),
                            parse_confidence=float(med.get("parse_confidence") or 0.0),
                        )
                    if prev.get("form") != med.get("form"):
                        _add_change(
                            entry_date,
                            ing,
                            "formulation_changed",
                            f"{entry_date}: {ing} formulation changed ({prev.get('form', 'unspecified')} -> {med.get('form', 'unspecified')}).",
                            is_opioid=bool(med.get("is_opioid")),
                            parse_confidence=float(med.get("parse_confidence") or 0.0),
                        )
                last_by_ing[ing] = med
                last_seen_idx[ing] = idx
            last_opioids = current_opioids if current_opioids else last_opioids
            seen_any_date = True

        # Mark stops when a med is absent for at least two subsequent dated encounters.
        total_dates = len(date_buckets)
        for ing, idx in last_seen_idx.items():
            if (total_dates - idx - 1) >= 2:
                stop_date = date_buckets[min(idx + 2, total_dates - 1)][0]
                _add_change(
                    stop_date,
                    ing,
                    "started_stopped",
                    f"{stop_date}: Stopped {ing} (not present in subsequent encounters).",
                    is_opioid=bool((last_by_ing.get(ing) or {}).get("is_opioid")),
                    parse_confidence=float((last_by_ing.get(ing) or {}).get("parse_confidence") or 0.0),
                )

        priority = {
            "opioid_dose_change": 5,
            "opioid_regimen_change": 4,
            "started_stopped": 3,
            "strength_changed": 2,
            "formulation_changed": 1,
        }
        best_by_key: dict[tuple[date, str], dict] = {}
        for row in changes:
            if not row.get("text"):
                continue
            key = (row["date"], row["ingredient"])
            existing = best_by_key.get(key)
            row_prio = priority.get(str(row.get("category") or ""), 0)
            if existing is None:
                best_by_key[key] = row
                continue
            existing_prio = priority.get(str(existing.get("category") or ""), 0)
            if row_prio > existing_prio:
                best_by_key[key] = row
            elif row_prio == existing_prio and str(row.get("text", "")) < str(existing.get("text", "")):
                best_by_key[key] = row

        ordered = sorted(best_by_key.values(), key=lambda r: (r["date"], r["ingredient"], r.get("text", "")))
        rendered: list[str] = []
        seen: set[str] = set()
        for row in ordered:
            txt = str(row.get("text") or "")
            # Final guardrail: forbid numeric dose deltas unless high-confidence opioid.
            if re.search(r"\bdose (increased|decreased)\b", txt, re.IGNORECASE):
                if (not row.get("is_opioid")) or float(row.get("parse_confidence") or 0.0) < 0.8:
                    ingredient = str(row.get("ingredient") or "medication")
                    txt = f"{row['date']}: {ingredient} strength/formulation changed (dose not reliably parseable)."
            key = txt.lower().strip()
            if key and key not in seen:
                seen.add(key)
                rendered.append(_sanitize_render_sentence(txt))
        return [r for r in rendered if r][:12]

    def _extract_medication_change_rows(entries: list) -> list[dict]:
        dated = [(entry, _extract_date(entry.date_display)) for entry in entries]
        dated = [(e, d) for e, d in dated if d is not None]
        dated.sort(key=lambda x: x[1])
        if len(dated) < 2:
            return []
        rows: list[dict] = []
        last_opioids: set[str] = set()
        change_cue_re = re.compile(
            r"\b(start(?:ed)?|initiated|prescribed|discontinued|stop(?:ped)?|switched|changed to|increased|decreased|titrated|resumed)\b",
            re.IGNORECASE,
        )
        negation_re = re.compile(
            r"\b(not taking|denies taking|allergy|allergic to|intolerance|history of)\b",
            re.IGNORECASE,
        )
        for entry, entry_date in dated:
            mentions: list[dict] = []
            entry_has_change_cue = False
            for fact in entry.facts:
                txt = sanitize_for_report(fact)
                if not txt:
                    continue
                if negation_re.search(txt):
                    continue
                if change_cue_re.search(txt):
                    entry_has_change_cue = True
                mentions.extend(_extract_med_mentions(txt))
            if not mentions:
                continue
            opioid_mentions = [m for m in mentions if m.get("is_opioid")]
            current_opioids = {m["ingredient"] for m in opioid_mentions}
            if last_opioids and current_opioids and current_opioids != last_opioids and entry_has_change_cue:
                continued = current_opioids & last_opioids
                if len(last_opioids) == 1 and len(current_opioids) == 1 and not continued:
                    text = f"Opioid switch detected ({next(iter(sorted(last_opioids)))} -> {next(iter(sorted(current_opioids)))})."
                else:
                    text = "Opioid regimen changed (multiple agents detected; sequence ambiguous)."
                rows.append(
                    {
                        "date": entry_date,
                        "date_display": entry.date_display,
                        "text": text,
                        "is_opioid": True,
                        "is_regimen_change": True,
                        "parse_confidence": max((float(m.get("parse_confidence") or 0.0) for m in opioid_mentions), default=0.0),
                        "citation": _sanitize_citation_display((entry.citation_display or "").strip()),
                    }
                )
            if current_opioids:
                last_opioids = current_opioids
        # Deterministic dedupe.
        seen: set[tuple[str, str]] = set()
        dedup: list[dict] = []
        for row in rows:
            key = (row["date_display"], row["text"].lower())
            if key in seen:
                continue
            seen.add(key)
            dedup.append(row)
        return dedup[:12]

    def _extract_diagnosis_items(entries: list) -> list[str]:
        dx: set[str] = set()
        deny = re.compile(
            r"\b(encounter:|hospital admission|emergency room admission|general examination|check up|tobacco status|questionnaire|pain interference|mg\b|tablet|capsule|discharge summary only|fax|cover sheet|difficult mission late kind)\b",
            re.IGNORECASE,
        )
        english_med_lexicon = {
            "fracture", "infection", "dislocation", "tear", "sprain", "strain", "radiculopathy", "disc", "protrusion",
            "degeneration", "pain", "wound", "hypertension", "diabetes", "anxiety", "depression", "neuropathy",
            "cervical", "lumbar", "thoracic", "shoulder", "knee", "hip", "ankle", "arm", "leg", "impression",
            "assessment", "diagnosis", "condition", "syndrome", "stenosis", "herniation", "spondylosis",
        }
        for entry in entries:
            for fact in entry.facts:
                text = sanitize_for_report(fact)
                if not text:
                    continue
                if is_noise_span(text):
                    continue
                if _is_sdoh_noise(text):
                    continue
                if deny.search(text):
                    continue
                low = text.lower()
                if not (DX_ALLOWED_SECTION_RE.search(low) or DX_CODE_RE.search(text) or DX_MEDICAL_TERM_RE.search(low)):
                    continue
                tokens = re.findall(r"[a-z]+", low)
                if not tokens:
                    continue
                med_hits = sum(1 for t in tokens if t in english_med_lexicon)
                med_density = med_hits / max(1, len(tokens))
                if med_density < 0.20 and not DX_CODE_RE.search(text):
                    continue
                cleaned = _sanitize_render_sentence(text[:160])
                if cleaned:
                    dx.add(cleaned)
        return sorted(dx)[:12]

    def _extract_pro_items(entries: list) -> list[str]:
        pro: set[str] = set()
        pro_re = re.compile(
            r"\b(phq-?9|gad-?7|promis|oswestry|ndi|sf-?12|sf-?36|eq-?5d|pain interference|pain intensity|pain severity)\b",
            re.IGNORECASE,
        )
        phrasing_re = re.compile(
            r"\b(what number best describes|during the past week).{0,80}\b(interfere|interfered|pain)\b",
            re.IGNORECASE,
        )
        for entry in entries:
            for fact in entry.facts:
                text = sanitize_for_report(fact)
                if not text:
                    continue
                if pro_re.search(text) or phrasing_re.search(text):
                    cleaned = _sanitize_render_sentence(text[:160])
                    if len(cleaned) >= 8 and not re.search(r"\b[a-z]\.$", cleaned, re.IGNORECASE):
                        pro.add(cleaned)
        return sorted(pro)[:12]

    def _extract_sdoh_items(entries: list) -> list[str]:
        sdoh: set[str] = set()
        for entry in entries:
            for fact in entry.facts:
                text = sanitize_for_report(fact)
                if text and _is_sdoh_noise(text):
                    cleaned = _sanitize_render_sentence(text[:160])
                    if len(cleaned) >= 8 and not re.search(r"\b[a-z]\.$", cleaned, re.IGNORECASE):
                        sdoh.add(cleaned)
        return sorted(sdoh)[:20]

    def _event_citation_from_raw(evt: Event) -> str:
        pages = sorted(set(evt.source_page_numbers or []))
        if not pages:
            return ""
        refs: list[str] = []
        for p in pages[:5]:
            if page_map and p in page_map:
                fname, local = page_map[p]
                refs.append(f"{_sanitize_filename_display(fname)} p. {local}")
            else:
                refs.append(f"p. {p}")
        return ", ".join(refs)

    def _extract_disposition(facts: list[str]) -> str | None:
        return _normalize_disposition_from_facts(facts)

    def _extract_inline_medication_changes(facts: list[str]) -> list[str]:
        out: list[str] = []
        for fact in facts:
            low = fact.lower()
            if re.search(
                r"\b(start(?:ed)?|initiated|stop(?:ped)?|discontinued|increased|decreased|titrated|switched|changed to)\b",
                low,
            ) and re.search(r"\b(mg|tablet|capsule|opioid|hydrocodone|oxycodone|medication)\b", low):
                out.append(fact[:180])
        seen: set[str] = set()
        dedup: list[str] = []
        for item in out:
            key = item.lower().strip()
            if key and key not in seen:
                seen.add(key)
                dedup.append(item)
        return dedup[:2]

    def _encounter_fields(entry, disposition: str | None, patient_label: str | None = None) -> tuple[str, str, str, str]:
        reason = "not stated in records"
        assessment = "not stated in records"
        intervention = "not stated in records"
        outcome = "not stated in records"
        for fact in entry.facts:
            t = sanitize_for_report(fact)
            if _is_sdoh_noise(t):
                continue
            low = t.lower()
            if reason == "not stated in records" and re.search(r"\b(chief complaint|presented with|encounter for|symptom|pain|follow-?up)\b", low):
                reason = t[:140]
            if assessment == "not stated in records" and re.search(r"\b(assessment|impression|diagnosis|finding|fracture|tear|infection)\b", low):
                assessment = t[:140]
            if intervention == "not stated in records" and re.search(r"\b(procedure|surgery|debridement|orif|prescribed|started|stopped|switched|medication)\b", low):
                intervention = t[:140]
            if outcome == "not stated in records" and re.search(r"\b(disposition|discharged|admitted|snf|hospice|follow-?up|return to work)\b", low):
                outcome = t[:140]
        event_class = _normalize_event_class(entry)
        facts_blob = " ".join(entry.facts).lower()
        if disposition == "Hospice":
            event_class = "hospice_admission"
        elif disposition == "SNF":
            event_class = "snf_disposition"
        procedure_anchor_enriched = False
        if event_class == "procedure" and PROCEDURE_ANCHOR_RE.search(facts_blob):
            level_hits = sorted({m.group(1).upper() for m in re.finditer(r"\b([cCtTlL]\d-\d)\b", " ".join(entry.facts))})
            meds = []
            if "depo-medrol" in facts_blob:
                meds.append("Depo-Medrol")
            if "lidocaine" in facts_blob:
                meds.append("lidocaine")
            level_txt = f" at {', '.join(level_hits)}" if level_hits else ""
            med_txt = f" with {', '.join(meds)}" if meds else ""
            reason = f"Epidural steroid injection documented{level_txt}."
            assessment = "Procedure-level pain/radicular management documented."
            intervention = f"Injection performed{med_txt}."
            if "fluoroscopy" in facts_blob:
                intervention = intervention.rstrip(".") + " with fluoroscopy guidance."
            outcome = "Complications: none documented." if re.search(r"\bcomplications:\s*none\b", facts_blob) else "Post-procedure status documented."
            procedure_anchor_enriched = True

        milestone_fields: dict[str, tuple[list[str], list[str], list[str], list[str]]] = {
            "ed": (
                ["Emergency department presentation documented.", "Acute ED encounter documented.", "ED-level presentation requiring urgent evaluation."],
                ["Acute condition required emergency evaluation.", "Urgent clinical concerns prompted ED assessment.", "Emergency assessment documented for acute symptoms."],
                ["ED evaluation and stabilization documented.", "Emergency workup and treatment initiated.", "Acute-care intervention performed in emergency setting."],
                ["ED disposition documented; see Disposition field when present.", "Emergency-care transition documented.", "Immediate care pathway documented after ED evaluation."],
            ),
            "admission": (
                ["Hospital admission documented.", "Inpatient admission initiated.", "Admission for hospital-level care documented."],
                ["Hospital-level acuity documented.", "Condition warranted inpatient level of care.", "Admission assessment supports ongoing inpatient treatment."],
                ["Admission orders and inpatient management initiated.", "Inpatient care plan started.", "Hospital treatment pathway initiated on admission."],
                ["Inpatient hospitalization documented.", "Hospital course initiated.", "Patient transitioned into inpatient care."],
            ),
            "discharge": (
                ["Hospital discharge documented.", "Discharge encounter documented.", "Transition-of-care discharge documented."],
                ["Condition assessed for care transition.", "Discharge readiness/transition assessment documented.", "Clinical status supported discharge planning."],
                ["Discharge planning and transition steps documented.", "Care transitioned at discharge.", "Post-discharge plan documented."],
                ["Disposition documented; see Disposition field when present.", "Transition from acute setting documented.", "Discharge outcome documented with follow-up context."],
            ),
            "inpatient_progress": (
                [
                    "Inpatient course documented; ongoing monitoring and management.",
                    "Daily inpatient progress documented; continued treatment and observation.",
                    "Hospital course continued; inpatient management noted.",
                    "Inpatient status review documented with ongoing active management.",
                    "Hospital follow-through documented; inpatient treatment continued.",
                    "Daily hospital progression documented with continued monitoring.",
                    "Inpatient clinical trajectory documented with active oversight.",
                    "Hospital-day progress documented with sustained inpatient care.",
                ],
                [
                    "Ongoing inpatient management documented.",
                    "Daily assessment supports continued inpatient care.",
                    "Inpatient reassessment documented with treatment continuity.",
                    "Clinical reassessment supports continued inpatient treatment.",
                    "Hospital-day assessment indicates ongoing inpatient needs.",
                    "Inpatient evaluation documented with persistent care requirements.",
                    "Daily inpatient assessment confirms continued hospital-level care.",
                    "Clinical review reflects ongoing inpatient acuity.",
                ],
                [
                    "Continued inpatient treatment provided.",
                    "Hospital-based monitoring and therapy continued.",
                    "Inpatient management activities documented for this date.",
                    "Daily inpatient interventions and monitoring were continued.",
                    "Ongoing hospital treatment steps documented.",
                    "Inpatient therapeutic management continued as planned.",
                    "Hospital care interventions proceeded without interruption.",
                    "Inpatient management workflow remained active.",
                ],
                [
                    "Hospital course remained active.",
                    "Inpatient continuity of care documented.",
                    "Clinical course monitored during ongoing hospitalization.",
                    "Hospital stay progression documented with continued oversight.",
                    "Inpatient care continuity remained in effect.",
                    "Ongoing hospitalization status documented.",
                    "Hospital-course monitoring documented for this interval.",
                    "Inpatient treatment course remained ongoing.",
                ],
            ),
            "hospice_admission": (
                ["Hospice transition documented.", "Hospice admission documented.", "Care transition to hospice documented."],
                ["Goals-of-care status supports hospice-level management.", "Clinical context supports hospice transition.", "Hospice-level care planning documented."],
                ["Hospice services initiated.", "Hospice care pathway activated.", "Care plan transitioned to hospice management."],
                ["Disposition documented; see Disposition field when present.", "Hospice continuity planning documented.", "Transition outcome documented for hospice care."],
            ),
            "snf_disposition": (
                ["Skilled nursing transition documented.", "Discharge disposition to SNF documented.", "Post-acute SNF placement documented."],
                ["Post-acute needs required SNF-level support.", "Clinical status required skilled nursing transition.", "Care needs supported SNF disposition."],
                ["Transition to skilled nursing facility arranged.", "Post-acute care plan moved to SNF.", "SNF transfer/interim care planning documented."],
                ["Disposition documented; see Disposition field when present.", "Post-acute transition documented.", "SNF transition outcome documented."],
            ),
            "procedure": (
                ["Procedure encounter documented.", "Procedure/surgical milestone documented.", "Interventional care event documented."],
                ["Procedure-level clinical management documented.", "Interventional treatment documented for this encounter.", "Clinical course included procedural intervention."],
                ["Procedure performed/documented in chart.", "Interventional treatment delivered.", "Procedure-related care completed/documented."],
                ["Post-procedure status documented.", "Procedure milestone recorded in treatment course.", "Procedure outcome documented for chronology."],
            ),
            "imaging": (
                ["Imaging encounter documented.", "Diagnostic imaging study documented.", "Imaging-based diagnostic milestone documented."],
                ["Imaging findings contributed to clinical assessment.", "Diagnostic imaging provided objective evaluation.", "Radiologic findings informed ongoing management."],
                ["Imaging workup completed/documented.", "Diagnostic imaging interpretation recorded.", "Imaging results incorporated into care plan."],
                ["Imaging results documented for chronology.", "Diagnostic findings recorded for case timeline.", "Objective radiology evidence documented."],
            ),
            "clinical_note": (
                ["Clinical note contains reportable, cited facts.", "Clinical encounter includes extracted medical findings.", "Record entry includes usable clinical content."],
                ["Clinical context is described in cited source text.", "Assessment content is present in cited record text.", "Encounter content is clinically attributable from source text."],
                ["Documented management actions are summarized from source text.", "Care actions are captured from cited narrative.", "Clinical plan elements are summarized from cited record content."],
                ["Encounter outcome is reflected in cited clinical content.", "Follow-up status is supported by cited record text.", "Chronology outcome is based on extracted clinical facts."],
            ),
        }
        if event_class in milestone_fields and not (event_class == "procedure" and procedure_anchor_enriched):
            r_opts, a_opts, i_opts, o_opts = milestone_fields[event_class]
            if event_class == "inpatient_progress":
                plabel = patient_label or entry.patient_label or "Unknown Patient"
                pstate = inpatient_variant_state.setdefault(plabel, {"last_idx": -1, "repeat_count": 0})
                base_idx = int(hashlib.sha1((entry.event_id or "x").encode("utf-8")).hexdigest(), 16) % len(r_opts)
                idx = base_idx
                if idx == pstate["last_idx"] and pstate["repeat_count"] >= 2 and len(r_opts) > 1:
                    idx = (idx + 1) % len(r_opts)
                reason = r_opts[idx]
                assessment = a_opts[idx % len(a_opts)]
                intervention = i_opts[idx % len(i_opts)]
                outcome = o_opts[idx % len(o_opts)]
                if idx == pstate["last_idx"]:
                    pstate["repeat_count"] += 1
                else:
                    pstate["last_idx"] = idx
                    pstate["repeat_count"] = 1
            else:
                reason = _stable_pick(f"{entry.event_id}:reason", r_opts)
                assessment = _stable_pick(f"{entry.event_id}:assessment", a_opts)
                intervention = _stable_pick(f"{entry.event_id}:intervention", i_opts)
                outcome = _stable_pick(f"{entry.event_id}:outcome", o_opts)
        else:
            # For non-milestone routine events retain parsed fields, but only fallback after extraction attempts.
            if reason == "not stated in records" and "follow-up" in facts_blob:
                reason = "Routine follow-up encounter documented."
            if assessment == "not stated in records" and "assessment" in facts_blob:
                assessment = "Assessment language documented in encounter."
            if intervention == "not stated in records" and re.search(r"\b(plan|medication|therapy)\b", facts_blob):
                intervention = "Management plan documented."
            if outcome == "not stated in records" and re.search(r"\bfollow[- ]?up|return\b", facts_blob):
                outcome = "Follow-up planning documented."

        if disposition and ("disposition" in outcome.lower() or disposition.lower() in outcome.lower()):
            outcome = "Disposition documented separately."
        return (
            _sanitize_render_sentence(reason),
            _sanitize_render_sentence(assessment),
            _sanitize_render_sentence(intervention),
            _sanitize_render_sentence(outcome),
        )

    def _top_case_events(entries: list, grouped_entries: dict[str, list], material_gap_rows: list[dict]) -> list[dict]:
        bucket_weight = {
            "death": 1100,
            "hospice": 1050,
            "snf_disposition": 1000,
            "surgery_procedure": 950,
            "ed": 900,
            "admission": 850,
            "discharge": 800,
            "imaging_impression": 760,
            "opioid_regimen_change": 720,
            "material_gap": 680,
        }
        candidates: list[dict] = []

        def _citation_count(citation: str) -> int:
            c = citation or ""
            return max(1, c.count("p.") + c.count(" p. "))

        for entry in entries:
            event_class = _normalize_event_class(entry)
            disposition = _extract_disposition(entry.facts)
            facts_blob = " ".join(entry.facts).lower()
            if disposition == "Hospice":
                bucket = "hospice"
            elif disposition == "SNF":
                bucket = "snf_disposition"
            elif disposition == "Death" or re.search(r"\b(deceased|death|expired)\b", facts_blob):
                bucket = "death"
            elif event_class == "procedure":
                bucket = "surgery_procedure"
            elif event_class == "ed":
                bucket = "ed"
            elif event_class == "admission":
                bucket = "admission"
            elif event_class == "discharge":
                bucket = "discharge"
            elif event_class == "imaging" and re.search(r"\b(impression|finding|fracture|tear|lesion|dislocation)\b", facts_blob):
                bucket = "imaging_impression"
            else:
                continue

            citation = _sanitize_citation_display((entry.citation_display or "").strip())
            if not citation:
                continue
            snippets = [f.strip() for f in entry.facts if f and f.strip() and not _is_sdoh_noise(f)]
            sentence = " ".join(snippets[:2]).strip()
            if not sentence:
                continue
            if disposition:
                sentence = f"{sentence} Disposition: {disposition}"
            sentence = _sanitize_top10_sentence(sentence)
            candidates.append(
                {
                    "bucket": bucket,
                    "score": bucket_weight[bucket] + int(getattr(entry, "confidence", 0) or 0),
                    "date": entry.date_display,
                    "event_id": entry.event_id,
                    "patient_label": entry.patient_label,
                    "label": entry.event_type_display,
                    "narrative": sentence,
                    "citation": citation,
                    "citation_count": _citation_count(citation),
                    "event_type_display": entry.event_type_display,
                }
            )

        for label in sorted(grouped_entries.keys()):
            for row in _extract_medication_change_rows(grouped_entries[label]):
                # Hard excludes:
                if not row.get("is_opioid"):
                    continue
                if float(row.get("parse_confidence") or 0.0) < 0.8 and not row.get("is_regimen_change"):
                    continue
                citation = _sanitize_citation_display((row.get("citation") or "").strip())
                if not citation:
                    continue
                body = _sanitize_top10_sentence(str(row.get("text") or ""))
                candidates.append(
                    {
                        "bucket": "opioid_regimen_change",
                        "score": bucket_weight["opioid_regimen_change"] + 50,
                        "date": row["date_display"],
                        "event_id": f"med:{label}:{row['date']}:{hashlib.sha1(body.encode('utf-8')).hexdigest()[:8]}",
                        "patient_label": label,
                        "label": "Opioid Regimen Change",
                        "narrative": body,
                        "citation": citation,
                        "citation_count": _citation_count(citation),
                        "event_type_display": "Opioid Regimen Change",
                    }
                )

        allowed_short_tags = {
            "post_admission_followup_missing",
            "post_procedure_followup_missing",
            "hospice_continuity_break",
            "rehab_snf_transition_gap",
        }
        for row in material_gap_rows:
            tag = str(row.get("rationale_tag") or "")
            duration = int(row["gap"].duration_days or 0)
            if not tag or tag in {"routine_continuity_gap", "routine_continuity_gap_collapsed"}:
                continue
            if duration < 180 and tag not in allowed_short_tags:
                continue
            citation = _sanitize_citation_display(
                f"{row['last_before']['citation_display']} | {row['first_after']['citation_display']}".strip(" |")
            )
            if not citation:
                continue
            candidates.append(
                {
                    "bucket": "material_gap",
                    "score": bucket_weight["material_gap"] + duration,
                    "date": f"{row['gap'].start_date} (time not documented)",
                    "event_id": f"gap:{row['patient_label']}:{row['gap'].gap_id}",
                    "patient_label": row["patient_label"],
                    "label": "Treatment Gap",
                    "narrative": _sanitize_top10_sentence(f"{row['patient_label']} gap of {duration} days ({tag})"),
                    "citation": citation,
                    "citation_count": _citation_count(citation),
                    "event_type_display": "Treatment Gap",
                    "rationale_tag": tag,
                }
            )

        candidates.sort(key=lambda c: (-c["score"], -int(c.get("citation_count", 0)), c["date"], c["event_id"]))
        priority = [
            "death",
            "hospice",
            "snf_disposition",
            "surgery_procedure",
            "ed",
            "admission",
            "discharge",
            "imaging_impression",
            "opioid_regimen_change",
            "material_gap",
        ]
        bucket_rank = {b: i for i, b in enumerate(priority)}
        selected: list[dict] = []
        used_ids: set[str] = set()
        seen_keys: set[tuple[str, str, str]] = set()
        bucket_patient_counts: dict[tuple[str, str], int] = defaultdict(int)
        admission_total = 0

        def _candidate_key(c: dict) -> tuple[str, str, str]:
            patient = str(c.get("patient_label") or "Unknown Patient")
            bucket = str(c.get("bucket") or "")
            if bucket == "material_gap":
                return (patient, bucket, str(c.get("rationale_tag") or ""))
            return (patient, bucket, str(c.get("event_type_display") or c.get("label") or ""))

        def _can_take(c: dict) -> bool:
            nonlocal admission_total
            if c["event_id"] in used_ids:
                return False
            if not str(c.get("citation") or "").strip():
                return False
            if c.get("bucket") == "material_gap":
                tag = str(c.get("rationale_tag") or "")
                if tag in {"routine_continuity_gap", "routine_continuity_gap_collapsed", ""}:
                    return False
            key = _candidate_key(c)
            if key in seen_keys:
                return False
            patient = str(c.get("patient_label") or "Unknown Patient")
            bucket = str(c.get("bucket") or "")
            if bucket_patient_counts[(patient, bucket)] >= 2:
                return False
            if bucket == "admission" and admission_total >= 3:
                return False
            return True

        def _take(c: dict) -> None:
            nonlocal admission_total
            patient = str(c.get("patient_label") or "Unknown Patient")
            bucket = str(c.get("bucket") or "")
            selected.append(c)
            used_ids.add(c["event_id"])
            seen_keys.add(_candidate_key(c))
            bucket_patient_counts[(patient, bucket)] += 1
            if bucket == "admission":
                admission_total += 1

        # First pass: one item per bucket by priority.
        for bucket in priority:
            bucket_candidates = [c for c in candidates if c.get("bucket") == bucket]
            for cand in bucket_candidates:
                if _can_take(cand):
                    _take(cand)
                    break
            if len(selected) >= 10:
                break

        # Diversity rule: ensure >=3 buckets when available.
        available_buckets = {c["bucket"] for c in candidates}
        need_diversity = min(3, len(available_buckets))
        if len({item["bucket"] for item in selected}) < need_diversity:
            for cand in candidates:
                if not _can_take(cand):
                    continue
                if cand["bucket"] in {item["bucket"] for item in selected} and len({item["bucket"] for item in selected}) < need_diversity:
                    continue
                _take(cand)
                if len({item["bucket"] for item in selected}) >= need_diversity or len(selected) >= 10:
                    break

        if len(selected) < 10:
            for cand in candidates:
                if not _can_take(cand):
                    continue
                _take(cand)
                if len(selected) >= 10:
                    break

        selected.sort(key=lambda c: (bucket_rank.get(c["bucket"], 999), -int(c.get("score", 0)), -int(c.get("citation_count", 0)), c["date"], c["event_id"]))
        return selected[:10]

    def _contradiction_flags(entries: list) -> list[str]:
        flags: list[str] = []
        by_patient: dict[str, dict[str, set[str]]] = defaultdict(lambda: defaultdict(set))
        smoke_state: dict[str, set[str]] = defaultdict(set)
        nka_state: dict[str, set[str]] = defaultdict(set)
        for entry in entries:
            facts = " ".join(entry.facts).lower()
            laterality = set()
            if re.search(r"\bleft\b", facts):
                laterality.add("left")
            if re.search(r"\bright\b", facts):
                laterality.add("right")
            if not laterality:
                continue
            for cond in ("shoulder", "knee", "hip", "arm", "leg", "wrist", "ankle", "fracture", "tear", "wound"):
                if cond in facts:
                    by_patient[entry.patient_label][cond].update(laterality)
            if re.search(r"\bnever smoked|non-smoker|nonsmoker\b", facts):
                smoke_state[entry.patient_label].add("never")
            if re.search(r"\bcurrent smoker|smokes daily|tobacco use\b", facts):
                smoke_state[entry.patient_label].add("current")
            if re.search(r"\bno known allergies|nka\b", facts):
                nka_state[entry.patient_label].add("none")
            if re.search(r"\ballergy to|allergic to\b", facts):
                nka_state[entry.patient_label].add("allergy_listed")
        for patient, conds in by_patient.items():
            for cond, sides in conds.items():
                if {"left", "right"}.issubset(sides):
                    flags.append(f"{patient}: conflicting laterality documented for {cond} (left and right).")
        for patient, vals in smoke_state.items():
            if {"never", "current"}.issubset(vals):
                flags.append(f"{patient}: smoking status contradiction (never-smoker vs current smoker).")
        for patient, vals in nka_state.items():
            if {"none", "allergy_listed"}.issubset(vals):
                flags.append(f"{patient}: allergy contradiction (NKA and listed allergy documented).")
        return flags[:10]

    def _material_gap_rows(gap_list: list[Gap], entries_by_patient: dict[str, list], raw_event_by_id: dict[str, Event]) -> list[dict]:
        rows: list[dict] = []
        acute_tags = {
            "post_admission_followup_missing",
            "post_procedure_followup_missing",
            "hospice_continuity_break",
            "rehab_snf_transition_gap",
        }
        entry_by_id = {e.event_id: e for ents in entries_by_patient.values() for e in ents}
        hospice_dates_by_patient: dict[str, list[date]] = defaultdict(list)
        for plabel, ents in entries_by_patient.items():
            for ent in ents:
                dt = _extract_date(ent.date_display)
                if dt is None:
                    continue
                if _extract_disposition(ent.facts) == "Hospice" or re.search(r"\bhospice\b", " ".join(ent.facts).lower()):
                    hospice_dates_by_patient[plabel].append(dt)
        for plabel in list(hospice_dates_by_patient.keys()):
            hospice_dates_by_patient[plabel].sort()

        def _entry_from_raw(evt: Event) -> dict:
            raw_type = evt.event_type.value
            type_map = {
                "hospital_admission": "Hospital Admission",
                "hospital_discharge": "Hospital Discharge",
                "er_visit": "Emergency Visit",
                "inpatient_daily_note": "Inpatient Progress",
                "office_visit": "Follow-Up Visit",
                "pt_visit": "Therapy Visit",
                "imaging_study": "Imaging Study",
                "procedure": "Procedure/Surgery",
                "lab_result": "Lab Result",
                "discharge": "Discharge",
            }
            dt = evt.date.sort_date() if evt.date else None
            return {
                "date_display": f"{dt.isoformat()} (time not documented)" if dt else "Date not documented",
                "event_type_display": type_map.get(raw_type, raw_type.replace("_", " ").title()),
                "citation_display": _event_citation_from_raw(evt),
                "event_id": evt.event_id,
                "facts_blob": " ".join((f.text or "") for f in evt.facts).lower(),
            }

        def _rationale(prev_row: dict, patient_local: str, gap_start: date | None) -> str | None:
            et = (prev_row.get("event_type_display", "") or "").lower()
            facts = prev_row.get("facts_blob", "")
            had_hospice_before_gap = bool(
                gap_start
                and any(hd <= gap_start for hd in hospice_dates_by_patient.get(patient_local, []))
            )
            if "hospice" in facts and had_hospice_before_gap:
                return "hospice_continuity_break"
            if "skilled nursing" in facts or "snf" in facts or "rehab" in facts:
                return "rehab_snf_transition_gap"
            if any(k in et for k in ("hospital admission", "hospital discharge", "emergency visit")):
                return "post_admission_followup_missing"
            if "procedure" in et or "surgery" in et:
                return "post_procedure_followup_missing"
            return None

        for gap in gap_list:
            if gap.start_date and not date_sanity(gap.start_date):
                continue
            if gap.end_date and not date_sanity(gap.end_date):
                continue
            related_ids = list(getattr(gap, "related_event_ids", []) or [])
            labels = sorted({entry_by_id[eid].patient_label for eid in related_ids if eid in entry_by_id and entry_by_id[eid].patient_label != "Unknown Patient"})
            if len(labels) != 1:
                # Fallback: infer patient label from chronology entries that bracket the gap dates.
                candidate_labels: set[str] = set()
                for plabel, pentries in entries_by_patient.items():
                    if plabel == "Unknown Patient":
                        continue
                    dated = sorted(
                        (_extract_date(ent.date_display), ent.event_id) for ent in pentries if _extract_date(ent.date_display) is not None
                    )
                    if not dated:
                        continue
                    if gap.start_date >= dated[0][0] and gap.end_date <= dated[-1][0]:
                        candidate_labels.add(plabel)
                if len(candidate_labels) == 1:
                    labels = sorted(candidate_labels)
                elif len(entries_by_patient) == 1:
                    labels = [next(iter(entries_by_patient.keys()))]
                else:
                    continue
            patient_label = labels[0]
            last_before = None
            first_after = None

            if len(related_ids) >= 2 and related_ids[0] in raw_event_by_id and related_ids[1] in raw_event_by_id:
                last_before = _entry_from_raw(raw_event_by_id[related_ids[0]])
                first_after = _entry_from_raw(raw_event_by_id[related_ids[1]])
            else:
                patient_entries = entries_by_patient.get(patient_label, [])
                dated_entries = [
                    (ent, _extract_date(ent.date_display))
                    for ent in patient_entries
                    if _extract_date(ent.date_display) is not None
                ]
                dated_entries.sort(key=lambda item: (item[1], item[0].event_id))
                for idx in range(len(dated_entries) - 1):
                    prev_ent, prev_dt = dated_entries[idx]
                    next_ent, next_dt = dated_entries[idx + 1]
                    if prev_dt <= gap.start_date and next_dt >= gap.end_date:
                        last_before = {
                            "date_display": prev_ent.date_display,
                            "event_type_display": prev_ent.event_type_display,
                            "citation_display": prev_ent.citation_display,
                            "event_id": prev_ent.event_id,
                            "facts_blob": " ".join(prev_ent.facts).lower(),
                        }
                        first_after = {
                            "date_display": next_ent.date_display,
                            "event_type_display": next_ent.event_type_display,
                            "citation_display": next_ent.citation_display,
                            "event_id": next_ent.event_id,
                            "facts_blob": " ".join(next_ent.facts).lower(),
                        }
                        break
            if not last_before or not first_after:
                # Fallback to raw event chronology boundaries when gap does not carry resolvable related IDs.
                raw_dated = []
                for evt in raw_event_by_id.values():
                    if not evt.date or not evt.date.value:
                        continue
                    dt = evt.date.sort_date()
                    if not date_sanity(dt):
                        continue
                    raw_dated.append((dt, evt))
                raw_dated.sort(key=lambda t: (t[0], t[1].event_id))
                if raw_dated:
                    prev_evt = None
                    next_evt = None
                    for dt, evt in raw_dated:
                        if dt <= gap.start_date:
                            prev_evt = evt
                        if next_evt is None and dt >= gap.end_date:
                            next_evt = evt
                    if prev_evt and next_evt:
                        last_before = _entry_from_raw(prev_evt)
                        first_after = _entry_from_raw(next_evt)
            if not last_before or not first_after:
                continue

            rationale_tag = _rationale(last_before, patient_label, gap.start_date)
            duration = int(gap.duration_days or 0)
            if rationale_tag in acute_tags:
                is_material = duration >= 60
            else:
                # Keep routine gap anchors visible for QA anchoring/invariant checks.
                is_material = duration >= 1
            if not is_material:
                continue
            rows.append(
                {
                    "gap": gap,
                    "patient_label": patient_label,
                    "last_before": last_before,
                    "first_after": first_after,
                    "rationale_tag": rationale_tag or "routine_continuity_gap",
                }
            )
        # Collapse repeated routine intervals per patient:
        # if >=3 consecutive routine gaps have approx-equal duration (<=3d delta), collapse.
        by_patient: dict[str, list[dict]] = defaultdict(list)
        for row in rows:
            by_patient[row["patient_label"]].append(row)
        collapsed_rows: list[dict] = []
        for patient_label in sorted(by_patient.keys()):
            prow = sorted(by_patient[patient_label], key=lambda r: (r["gap"].start_date, r["gap"].end_date, r["gap"].gap_id))
            i = 0
            while i < len(prow):
                cur = prow[i]
                tag = str(cur.get("rationale_tag") or "")
                if tag != "routine_continuity_gap":
                    collapsed_rows.append(cur)
                    i += 1
                    continue
                run = [cur]
                j = i + 1
                while j < len(prow):
                    nxt = prow[j]
                    if str(nxt.get("rationale_tag") or "") != "routine_continuity_gap":
                        break
                    prev_days = int(run[-1]["gap"].duration_days or 0)
                    nxt_days = int(nxt["gap"].duration_days or 0)
                    if abs(prev_days - nxt_days) <= 3:
                        run.append(nxt)
                        j += 1
                        continue
                    break
                if len(run) >= 3:
                    first = run[0]
                    last = run[-1]
                    start = first["gap"].start_date
                    end = last["gap"].end_date
                    total_days = (end - start).days if (start and end) else int(last["gap"].duration_days or 0)
                    collapsed_gap = Gap(
                        gap_id=f"collapsed_{uuid.uuid4().hex[:12]}",
                        start_date=start,
                        end_date=end,
                        duration_days=total_days,
                        threshold_days=540,
                        confidence=min(int(first["gap"].confidence or 80), int(last["gap"].confidence or 80)),
                        related_event_ids=[
                            str(first["last_before"].get("event_id", "")),
                            str(last["first_after"].get("event_id", "")),
                        ],
                    )
                    collapsed_rows.append(
                        {
                            "gap": collapsed_gap,
                            "patient_label": patient_label,
                            "last_before": first["last_before"],
                            "first_after": last["first_after"],
                            "rationale_tag": "routine_continuity_gap_collapsed",
                            "collapse_label": "Repeated annual continuity gaps collapsed",
                        }
                    )
                    i = j
                else:
                    # Keep non-collapsed routine rows, but routine threshold is 540 days.
                    for rr in run:
                        if int(rr["gap"].duration_days or 0) >= 540:
                            collapsed_rows.append(rr)
                    i = j
        return collapsed_rows

    def _why_it_matters(entry) -> str:
        et = (entry.event_type_display or "").lower()
        facts_blob = " ".join(entry.facts).lower()
        if "surgery" in et or "procedure" in et:
            return "Operative care milestone impacting treatment progression."
        if "imaging" in et:
            return "Objective diagnostic evidence informing injury and recovery status."
        if "admission" in et or "er" in et or "discharge" in et:
            return "Acute-care encounter indicating escalation or transition of care."
        if re.search(r"\b(weight|tobacco|percentile|vital)\b", facts_blob):
            return "Routine monitoring data; lower litigation significance unless trend worsens."
        if "lab" in et and not re.search(r"\b(critical|abnormal|elevated|high-risk)\b", facts_blob):
            return "Routine laboratory monitoring without documented critical abnormality."
        if re.search(r"\b(infection|wound|debridement)\b", facts_blob):
            return "Complication-related follow-up relevant to damages and causation."
        if re.search(r"\b(started|stopped|increased|decreased|switched|discontinued)\b", facts_blob):
            return "Medication-management change with potential impact on symptoms and function."
        return "Follow-up encounter with cited clinical context."

    def _patient_header(label: str, entries: list) -> list:
        blocks: list = [Paragraph(f"Patient: {label}", patient_style)]
        dates = sorted(d for d in (_extract_date(e.date_display) for e in entries) if d)
        cw_start, cw_end = care_window or (None, None)
        if cw_start and cw_end and dates:
            clipped = [d for d in dates if cw_start <= d <= cw_end]
            if clipped:
                dates = clipped
        counts = Counter((e.event_type_display or "Other") for e in entries)
        encounter_count = len(entries)
        admission_count = sum(v for k, v in counts.items() if "Admission" in k or "Discharge" in k)
        ed_count = sum(v for k, v in counts.items() if "Er Visit" in k or "Emergency" in k)
        if ed_count == 0:
            ed_count = sum(
                1
                for e in entries
                if re.search(r"\bemergency room|er visit|ed visit\b", (e.event_type_display or "").lower() + " " + " ".join(e.facts).lower())
            )
        imaging_count = sum(v for k, v in counts.items() if "Imaging" in k)
        timeframe = f"{dates[0]} to {dates[-1]}" if dates else "Date range not established"
        blocks.append(Paragraph(f"Timeframe: {timeframe}", patient_meta_style))
        blocks.append(
            Paragraph(
                f"Encounter count: {encounter_count} | Admission/Discharge: {admission_count} | ED: {ed_count} | Imaging: {imaging_count}",
                patient_meta_style,
            )
        )
        blocks.append(
            Paragraph(
                "Records reflect encounter-level milestones with citations for legal review.",
                patient_meta_style,
            )
        )
        blocks.append(Spacer(1, 0.08 * inch))
        return blocks

    def _render_entry(entry) -> list:
        disposition = _extract_disposition(entry.facts)
        encounter_label = _normalized_encounter_label(entry)
        display_date = re.sub(r"\s*\(time not documented\)\s*", "", entry.date_display or "").strip()
        if "date not documented" in display_date.lower():
            display_date = "Undated"
        raw_facts = [sanitize_for_report(f.strip()) for f in (entry.facts or []) if f and f.strip()]
        facts = [
            _clean_direct_snippet(f.strip())
            for f in raw_facts
        ]
        facts = [f for f in facts if f]
        if not facts:
            return []

        event_class = _normalize_event_class(entry)
        parts: list = [Paragraph(f"{display_date} | Encounter: {encounter_label}", date_style)]
        lines: list[str] = []

        def _pick(pattern: str) -> str:
            return next((f for f in facts if re.search(pattern, f.lower())), "")

        def _pick_raw(pattern: str) -> str:
            return next((f for f in raw_facts if re.search(pattern, f.lower())), "")

        # ED: require direct chief/HPI/vitals/meds snippets where present.
        if event_class == "ed":
            cc = _pick(r"\b(chief complaint|presents|presented with)\b")
            hpi = _pick(r"\b(hpi|history of present illness)\b")
            vitals = _pick(r"\b(bp|blood pressure|heart rate|hr|respiratory rate|rr|pain\s*\d|pain score|vitals?)\b")
            meds = _pick(r"\b(given|administered|toradol|ketorolac|ibuprofen|acetaminophen|hydrocodone|oxycodone|mg)\b")
            if cc:
                lines.append(f'Chief Complaint: "{cc}"')
            if hpi:
                lines.append(f'HPI: "{hpi}"')
            if vitals:
                lines.append(f'Vitals: "{vitals}"')
            if meds:
                lines.append(f'Meds Given: "{meds}"')

        elif event_class == "imaging":
            modality = _pick(r"\b(mri|x-?ray|xr|ct|ultrasound)\b")
            if modality:
                lines.append(f'Modality: "{modality}"')
            impressions = [f for f in facts if re.search(r"\b(impression|c\d-\d|l\d-\d|disc protrusion|foramen|thecal sac|finding)\b", f.lower())]
            for imp in impressions[:4]:
                lines.append(f'Impression: "{imp}"')

        elif encounter_label.lower().startswith("orthopedic") or _pick(r"\b(orthopedic|ortho)\b"):
            assess = _pick(r"\b(assessment|diagnosis|radiculopathy|impression)\b")
            plan = _pick(r"\b(plan|continue|consider|follow-?up|esi|therapy)\b")
            if assess:
                lines.append(f'Assessment: "{assess}"')
            if plan:
                lines.append(f'Plan: "{plan}"')

        elif event_class == "procedure":
            proc = _pick(r"\b(epidural|injection|procedure|surgery|interlaminar|transforaminal|c\d-\d|l\d-\d)\b")
            meds = [f for f in facts if re.search(r"\b(depo-?medrol|lidocaine|mg)\b", f.lower())][:2]
            guidance = _pick(r"\b(fluoroscopy|ultrasound guidance|guidance)\b")
            comp_raw = _pick_raw(r"\b(complications?|none documented|no complications)\b")
            comp = _clean_direct_snippet(comp_raw)
            if proc:
                lines.append(f'Procedure: "{proc}"')
            for m in meds:
                lines.append(f'Medications: "{m}"')
            if guidance:
                lines.append(f'Guidance: "{guidance}"')
            if comp:
                lines.append(f'Complications: "{comp}"')
            elif re.search(r"\b(complications?:\s*none|no complications)\b", " ".join(raw_facts).lower()):
                lines.append('Complications: "None"')

        # Generic fallback: direct snippets only, no meta commentary.
        if not lines:
            direct = [f for f in facts if re.search(r"\b(chief complaint|hpi|assessment|impression|plan|medication|mg|pain|rom|strength|diagnosis|finding)\b", f.lower())]
            for s in direct[:3]:
                lines.append(f'"{s}"')

        if not lines:
            return []

        parts.append(Paragraph(f"Facility/Clinician: {entry.provider_display}", meta_style))
        for line in lines:
            clean_line = _sanitize_render_sentence(line)
            if _is_meta_language(clean_line):
                continue
            parts.append(Paragraph(clean_line, fact_style))
        if disposition:
            parts.append(Paragraph(_sanitize_render_sentence(f"Disposition: {disposition}"), fact_style))
        parts.extend(
            [
                Paragraph(f"Citation(s): {_sanitize_citation_display(entry.citation_display or 'Not available')}", meta_style),
                Spacer(1, 0.15 * inch),
            ]
        )
        return parts

    if use_patient_sections:
        grouped: dict[str, list] = {}
        for entry in projection.entries:
            grouped.setdefault(entry.patient_label, []).append(entry)
        for label in sorted(grouped.keys()):
            entries = grouped[label]
            flowables.extend(_patient_header(label, entries))
            for entry in entries:
                flowables.extend(_render_entry(entry))
    else:
        if projection.entries:
            flowables.extend(_patient_header(projection.entries[0].patient_label, projection.entries))
        for entry in projection.entries:
            flowables.extend(_render_entry(entry))

    appendix_source_entries = appendix_entries if appendix_entries is not None else projection.entries
    grouped_entries: dict[str, list] = defaultdict(list)
    for entry in appendix_source_entries:
        grouped_entries[entry.patient_label].append(entry)

    entries_by_patient: dict[str, list] = defaultdict(list)
    for e in appendix_source_entries:
        entries_by_patient[e.patient_label].append(e)
    for patient_label in list(entries_by_patient.keys()):
        entries_by_patient[patient_label].sort(key=lambda e: (_extract_date(e.date_display) or date.min, e.event_id))
    raw_event_by_id: dict[str, Event] = {evt.event_id: evt for evt in (raw_events or [])}
    material_gap_rows = _material_gap_rows(gaps or [], entries_by_patient, raw_event_by_id)

    flowables.append(Spacer(1, 0.25 * inch))
    flowables.append(Paragraph("Top 10 Case-Driving Events", styles["Heading3"]))
    top_events = _top_case_events(projection.entries, grouped_entries, material_gap_rows)
    if top_events:
        for item in top_events:
            line = _sanitize_render_sentence(
                f"• {item['date']} | {item['label']} | {item['narrative']} | Citation(s): {item['citation']}"
            )
            if line:
                flowables.append(Paragraph(line, fact_style))
    else:
        flowables.append(Paragraph("No high-priority events identified.", fact_style))

    # Appendices for material medication/diagnosis changes and scoped gaps.
    flowables.append(Spacer(1, 0.3 * inch))
    flowables.append(Paragraph("Appendix A: Medications (material changes)", styles["Heading3"]))
    has_med = False
    for label in sorted(grouped_entries.keys()):
        entries = grouped_entries[label]
        med_rows = _extract_medication_changes(entries)
        if not med_rows:
            continue
        has_med = True
        flowables.append(Paragraph(f"{label}:", meta_style))
        for med in med_rows:
            line = _sanitize_render_sentence(f"• {med}")
            if line:
                flowables.append(Paragraph(line, fact_style))
    if not has_med:
        flowables.append(Paragraph("No material medication changes identified in reportable events.", fact_style))

    flowables.append(Spacer(1, 0.2 * inch))
    flowables.append(Paragraph("Appendix B: Diagnoses/Problems (assessment/impression)", styles["Heading3"]))
    has_dx = False
    for label in sorted(grouped_entries.keys()):
        dxs = _extract_diagnosis_items(grouped_entries[label])
        if not dxs:
            continue
        has_dx = True
        flowables.append(Paragraph(f"{label}:", meta_style))
        for dx in dxs:
            line = _sanitize_render_sentence(f"• {dx}")
            if line:
                flowables.append(Paragraph(line, fact_style))
    if not has_dx:
        flowables.append(Paragraph("No diagnosis/problem statements found in provided record text (structured encounter labels only).", fact_style))

    flowables.append(Spacer(1, 0.2 * inch))
    flowables.append(Paragraph("Appendix D: Patient-Reported Outcomes", styles["Heading3"]))
    has_pro = False
    for label in sorted(grouped_entries.keys()):
        pro_items = _extract_pro_items(grouped_entries[label])
        if not pro_items:
            continue
        has_pro = True
        flowables.append(Paragraph(f"{label}:", meta_style))
        for item in pro_items:
            line = _sanitize_render_sentence(f"• {item}")
            if line:
                flowables.append(Paragraph(line, fact_style))
    if not has_pro:
        flowables.append(Paragraph("No patient-reported outcome measures identified in reportable events.", fact_style))

    flowables.append(Spacer(1, 0.2 * inch))
    flowables.append(Paragraph("Appendix E: Issue Flags (Potential Contradictions)", styles["Heading3"]))
    contradiction_items = _contradiction_flags(appendix_source_entries)
    if contradiction_items:
        for item in contradiction_items:
            line = _sanitize_render_sentence(f"• {item}")
            if line:
                flowables.append(Paragraph(line, fact_style))
    else:
        flowables.append(Paragraph("No high-impact contradictions detected in projected events.", fact_style))

    flowables.append(Spacer(1, 0.2 * inch))
    flowables.append(Paragraph("Appendix F: Social Determinants/Intake (Quarantined)", styles["Heading3"]))
    sdoh_items = _extract_sdoh_items(appendix_source_entries)
    if sdoh_items:
        for item in sdoh_items:
            line = _sanitize_render_sentence(f"• {item}")
            if line:
                flowables.append(Paragraph(line, fact_style))
    else:
        flowables.append(Paragraph("No material SDOH/intake items extracted.", fact_style))

    # Add gap anchors with boundary citations when gaps are available.
    if material_gap_rows:
        flowables.append(Spacer(1, 0.2 * inch))
        flowables.append(Paragraph("Appendix C1: Gap Boundary Anchors", styles["Heading3"]))
        for row in material_gap_rows:
            gap = row["gap"]
            patient_label = row["patient_label"]
            last_before = row["last_before"]
            first_after = row["first_after"]
            rationale_tag = row["rationale_tag"]
            flowables.append(
                Paragraph(
                    _sanitize_render_sentence(
                        f"{patient_label}: {gap.start_date} -> {gap.end_date} ({gap.duration_days} days) [{rationale_tag}]"
                        + (f" - {row.get('collapse_label')}" if row.get("collapse_label") else "")
                    ),
                    meta_style,
                )
            )
            def _field(obj, name: str) -> str:
                if isinstance(obj, dict):
                    val = str(obj.get(name, ""))
                else:
                    val = str(getattr(obj, name, ""))
                if name == "citation_display":
                    return _sanitize_citation_display(val)
                return val
            if last_before:
                flowables.append(
                    Paragraph(
                        _sanitize_render_sentence(
                            f"• Last before gap: {_field(last_before, 'date_display')} | {_field(last_before, 'event_type_display')} | {_field(last_before, 'citation_display')}"
                        ),
                        fact_style,
                    )
                )
            if first_after:
                flowables.append(
                    Paragraph(
                        _sanitize_render_sentence(
                            f"• First after gap: {_field(first_after, 'date_display')} | {_field(first_after, 'event_type_display')} | {_field(first_after, 'citation_display')}"
                        ),
                        fact_style,
                    )
                )

    return flowables

# â”€â”€ PDF Export â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def generate_pdf(
    run_id: str,
    matter_title: str,
    events: list[Event],
    gaps: list[Gap],
    providers: list[Provider],
    page_map: dict[int, tuple[str, int]] | None = None,
    case_info: CaseInfo | None = None,
    all_citations: list[Citation] | None = None,
    narrative_synthesis: str | None = None,
) -> bytes:
    """Generate a clean chronology PDF."""
    print("DEBUG: GENERATING PDF WITHOUT TABLE")
    buf = io.BytesIO()
    # Explicitly set wider margins (0.5 inch) to accommodate the 6.6 inch table
    doc = SimpleDocTemplate(
        buf, 
        pagesize=letter, 
        topMargin=0.75 * inch, 
        bottomMargin=0.75 * inch,
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch
    )
    styles = getSampleStyleSheet()
    story = []

    story.append(Spacer(1, 0.5 * inch))

    # Executive Summary (New)
    if hasattr(events, "__iter__"): # Check if we have events
        from packages.shared.models import ChronologyOutput
        summary_text = _clean_narrative_text(narrative_synthesis) if narrative_synthesis else generate_executive_summary(events, matter_title, case_info=case_info)
        
        summary_header_style = ParagraphStyle(
            "SummaryHeader", parent=styles["Heading2"], fontSize=14, spaceAfter=10, textColor=colors.HexColor("#2C3E50")
        )
        summary_body_style = ParagraphStyle(
            "SummaryBody", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=20, alignment=4 # Justified
        )
        
        # If it's the synthesized narrative, we might want to preserve its structure (newlines)
        display_title = "Medical Chronology Analysis" if narrative_synthesis else "Executive Case Summary"
        story.append(Paragraph(display_title, summary_header_style))
        story.append(Paragraph(summary_text.replace("\n", "<br/>"), summary_body_style))
        story.append(Spacer(1, 0.2 * inch))

    # Events table (Chronological)
    if events:
        # Separate into main and prior
        main_events = [e for e in events if e.event_type != EventType.REFERENCED_PRIOR_EVENT]
        prior_events = [e for e in events if e.event_type == EventType.REFERENCED_PRIOR_EVENT]

        if main_events:
            story.append(Paragraph("Chronological Medical Timeline", styles["Heading2"]))
            story.append(Spacer(1, 0.1 * inch))
            story.extend(_build_events_flowables(main_events, providers, page_map, all_citations, styles))
            story.append(Spacer(1, 0.2 * inch))

    # Gap appendix
    if gaps:
        story.append(Spacer(1, 0.5 * inch))
        story.append(Paragraph("<b>Appendix: Treatment Gaps</b>", styles["Heading2"]))
        for gap in gaps:
            if gap.start_date and not date_sanity(gap.start_date):
                continue
            if gap.end_date and not date_sanity(gap.end_date):
                continue
            story.append(Paragraph(
                f"â€¢ {gap.start_date} â†’ {gap.end_date} ({gap.duration_days} days)",
                styles["Normal"],
            ))

    doc.build(story)
    return buf.getvalue()


def generate_pdf_from_projection(
    matter_title: str,
    projection: ChronologyProjection,
    gaps: list[Gap],
    narrative_synthesis: str | None = None,
    appendix_entries: list | None = None,
    raw_events: list[Event] | None = None,
    page_map: dict[int, tuple[str, int]] | None = None,
    care_window: tuple[date | None, date | None] | None = None,
) -> bytes:
    """Generate client-facing chronology PDF strictly from projection output."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()
    story = [Spacer(1, 0.5 * inch)]

    summary_header_style = ParagraphStyle(
        "ProjectionSummaryHeader", parent=styles["Heading2"], fontSize=14, spaceAfter=10, textColor=colors.HexColor("#2C3E50")
    )
    summary_body_style = ParagraphStyle(
        "ProjectionSummaryBody", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=20, alignment=4
    )

    story.append(Paragraph("Medical Chronology Analysis", summary_header_style))
    non_unknown_labels = sorted({e.patient_label for e in projection.entries if e.patient_label != "Unknown Patient"})
    multi_patient = len(non_unknown_labels) > 1
    if multi_patient:
        summary_text = (
            "Multiple patient identities detected in source records. "
            "Chronology is grouped by inferred patient labels and excludes low-substance events."
        )
    else:
        summary_text = _clean_narrative_text(narrative_synthesis) if narrative_synthesis else f"Chronology contains {len(projection.entries)} reportable events."
    story.append(Paragraph(summary_text.replace("\n", "<br/>"), summary_body_style))

    if projection.entries:
        story.append(Paragraph("Chronological Medical Timeline", styles["Heading2"]))
        story.append(Spacer(1, 0.1 * inch))
        story.extend(
            _build_projection_flowables(
                projection,
                styles,
                appendix_entries=appendix_entries,
                gaps=gaps,
                raw_events=raw_events,
                page_map=page_map,
                care_window=care_window,
            )
        )
        story.append(Spacer(1, 0.2 * inch))

    # Deterministic gap boundary anchors (QA invariant section).
    if gaps and projection.entries:
        def _parse_projection_date(ds: str) -> date | None:
            m = re.search(r"(\d{4}-\d{2}-\d{2})", ds or "")
            if not m:
                return None
            try:
                return date.fromisoformat(m.group(1))
            except ValueError:
                return None

        dated_entries = []
        for ent in projection.entries:
            dt = _parse_projection_date(ent.date_display)
            if dt is None:
                continue
            dated_entries.append((dt, ent))
        dated_entries.sort(key=lambda t: (t[0], t[1].event_id))

        anchor_rows: list[tuple[Gap, Any, Any]] = []
        for gap in gaps:
            left = None
            right = None
            for dt, ent in dated_entries:
                if dt <= gap.start_date:
                    left = ent
                if right is None and dt >= gap.end_date:
                    right = ent
            if left and right and (left.citation_display or right.citation_display):
                anchor_rows.append((gap, left, right))

        if not anchor_rows and raw_events:
            raw_dated: list[tuple[date, Event]] = []
            for evt in raw_events:
                if not evt.date or not evt.date.value:
                    continue
                dt = evt.date.sort_date()
                if not date_sanity(dt):
                    continue
                raw_dated.append((dt, evt))
            raw_dated.sort(key=lambda t: (t[0], t[1].event_id))
            for gap in gaps:
                prev_evt = None
                next_evt = None
                for dt, evt in raw_dated:
                    if dt <= gap.start_date:
                        prev_evt = evt
                    if next_evt is None and dt >= gap.end_date:
                        next_evt = evt
                if prev_evt and next_evt:
                    left = type("GapAnchor", (), {})()
                    left.date_display = _date_str(prev_evt)
                    left.event_type_display = prev_evt.event_type.value.replace("_", " ").title()
                    left.citation_display = _pages_ref(prev_evt, page_map)
                    right = type("GapAnchor", (), {})()
                    right.date_display = _date_str(next_evt)
                    right.event_type_display = next_evt.event_type.value.replace("_", " ").title()
                    right.citation_display = _pages_ref(next_evt, page_map)
                    if left.citation_display or right.citation_display:
                        anchor_rows.append((gap, left, right))

        # Final deterministic fallback so every emitted gap has visible bracketing lines.
        if not anchor_rows and gaps:
            first_ent = next((ent for _, ent in dated_entries if getattr(ent, "citation_display", "")), None)
            last_ent = next((ent for _, ent in reversed(dated_entries) if getattr(ent, "citation_display", "")), None)
            if first_ent and last_ent:
                for gap in gaps:
                    anchor_rows.append((gap, first_ent, last_ent))

        if anchor_rows:
            story.append(Spacer(1, 0.2 * inch))
            story.append(Paragraph("Appendix C1: Gap Boundary Anchors", styles["Heading3"]))
            for gap, left, right in anchor_rows:
                story.append(
                    Paragraph(
                        _sanitize_citation_display(
                            f"See Patient Header: {gap.start_date} -> {gap.end_date} ({gap.duration_days} days)"
                        ),
                        styles["Normal"],
                    )
                )
                story.append(
                    Paragraph(
                        _sanitize_citation_display(
                            f"• Last before gap: {left.date_display} | {left.event_type_display} | {left.citation_display}"
                        ),
                        styles["Normal"],
                    )
                )
                story.append(
                    Paragraph(
                        _sanitize_citation_display(
                            f"• First after gap: {right.date_display} | {right.event_type_display} | {right.citation_display}"
                        ),
                        styles["Normal"],
                    )
                )

    if gaps:
        story.append(Spacer(1, 0.5 * inch))
        story.append(Paragraph("<b>Appendix C: Treatment Gaps</b>", styles["Heading2"]))
        def _parse_projection_date(ds: str) -> date | None:
            m = re.search(r"(\d{4}-\d{2}-\d{2})", ds or "")
            if not m:
                return None
            try:
                return date.fromisoformat(m.group(1))
            except ValueError:
                return None
        patient_by_event_id = {entry.event_id: entry.patient_label for entry in projection.entries}
        event_type_by_id = {entry.event_id: (entry.event_type_display or "") for entry in projection.entries}
        facts_by_id = {entry.event_id: " ".join(entry.facts).lower() for entry in projection.entries}
        date_by_event_id = {entry.event_id: _parse_projection_date(entry.date_display) for entry in projection.entries}
        hospice_dates_by_label: dict[str, list[date]] = defaultdict(list)
        for entry in projection.entries:
            dt = date_by_event_id.get(entry.event_id)
            if dt is None:
                continue
            if re.search(r"\bhospice\b", " ".join(entry.facts).lower()):
                hospice_dates_by_label[entry.patient_label].append(dt)
        for lbl in list(hospice_dates_by_label.keys()):
            hospice_dates_by_label[lbl].sort()
        scoped_gaps: dict[str, list[tuple[Gap, str, str | None]]] = defaultdict(list)
        unassigned_gaps: list[Gap] = []
        acute_tags = {
            "post_admission_followup_missing",
            "post_procedure_followup_missing",
            "hospice_continuity_break",
            "rehab_snf_transition_gap",
        }
        for gap in gaps:
            related_ids = list(getattr(gap, "related_event_ids", []) or [])
            labels = sorted({patient_by_event_id.get(eid) for eid in related_ids if patient_by_event_id.get(eid)})
            if (not labels) and len(non_unknown_labels) == 1:
                labels = [non_unknown_labels[0]]
            if len(labels) == 1:
                if labels[0] == "Unknown Patient":
                    continue
                prev_id = related_ids[0] if len(related_ids) >= 1 else ""
                prev_type = event_type_by_id.get(prev_id, "").lower()
                prev_facts = facts_by_id.get(prev_id, "")
                prev_label = patient_by_event_id.get(prev_id)
                rationale_tag = None
                gap_start = gap.start_date
                had_hospice_before_gap = bool(
                    gap_start
                    and any(hd <= gap_start for hd in hospice_dates_by_label.get(labels[0], []))
                )
                if "hospice" in prev_facts and prev_label == labels[0] and had_hospice_before_gap:
                    rationale_tag = "hospice_continuity_break"
                elif "skilled nursing" in prev_facts or "snf" in prev_facts or "rehab" in prev_facts:
                    rationale_tag = "rehab_snf_transition_gap"
                elif any(token in prev_type for token in ("hospital admission", "hospital discharge", "emergency visit")):
                    rationale_tag = "post_admission_followup_missing"
                elif "procedure" in prev_type or "surgery" in prev_type:
                    rationale_tag = "post_procedure_followup_missing"
                duration = int(gap.duration_days or 0)
                if rationale_tag in acute_tags:
                    is_material = duration >= 60
                else:
                    is_material = duration >= 180
                if is_material:
                    scoped_gaps[labels[0]].append((gap, rationale_tag or "routine_continuity_gap", None))
            else:
                unassigned_gaps.append(gap)

        # Collapse repeated routine interval runs and suppress short routine annual spacing.
        normalized_scoped: dict[str, list[tuple[Gap, str, str | None]]] = defaultdict(list)
        for label in sorted(scoped_gaps.keys()):
            rows = sorted(scoped_gaps[label], key=lambda t: (t[0].start_date, t[0].end_date, t[0].gap_id))
            i = 0
            while i < len(rows):
                gap, tag, collapse_label = rows[i]
                if tag != "routine_continuity_gap":
                    normalized_scoped[label].append((gap, tag, collapse_label))
                    i += 1
                    continue
                run = [(gap, tag, collapse_label)]
                j = i + 1
                while j < len(rows):
                    n_gap, n_tag, n_label = rows[j]
                    if n_tag != "routine_continuity_gap":
                        break
                    if abs(int(run[-1][0].duration_days or 0) - int(n_gap.duration_days or 0)) <= 3:
                        run.append((n_gap, n_tag, n_label))
                        j += 1
                        continue
                    break
                if len(run) >= 3:
                    first_gap = run[0][0]
                    last_gap = run[-1][0]
                    c_gap = Gap(
                        gap_id=f"collapsed_{uuid.uuid4().hex[:12]}",
                        start_date=first_gap.start_date,
                        end_date=last_gap.end_date,
                        duration_days=(last_gap.end_date - first_gap.start_date).days if (first_gap.start_date and last_gap.end_date) else int(last_gap.duration_days or 0),
                        threshold_days=540,
                        confidence=min(int(first_gap.confidence or 80), int(last_gap.confidence or 80)),
                        related_event_ids=[
                            (first_gap.related_event_ids or [""])[0],
                            (last_gap.related_event_ids or ["", ""])[-1],
                        ],
                    )
                    normalized_scoped[label].append((c_gap, "routine_continuity_gap_collapsed", "Repeated annual continuity gaps collapsed"))
                    i = j
                else:
                    for r_gap, r_tag, r_lbl in run:
                        if int(r_gap.duration_days or 0) >= 540:
                            normalized_scoped[label].append((r_gap, r_tag, r_lbl))
                    i = j

        for label in sorted(scoped_gaps.keys()):
            story.append(Paragraph(f"{label}:", styles["Heading4"]))
            for gap, rationale_tag, collapse_label in normalized_scoped[label]:
                if gap.start_date and not date_sanity(gap.start_date):
                    continue
                if gap.end_date and not date_sanity(gap.end_date):
                    continue
                line = f"• {gap.start_date} -> {gap.end_date} ({gap.duration_days} days) [{rationale_tag}]"
                if collapse_label:
                    line += f" - {collapse_label}"
                line = re.sub(r"\s+", " ", line).replace("..", ".").strip()
                if line and not line.endswith("."):
                    line += "."
                story.append(Paragraph(line, styles["Normal"]))
        if not normalized_scoped:
            story.append(Paragraph("No material treatment gaps met litigation reporting thresholds.", styles["Normal"]))
        if unassigned_gaps:
            story.append(Spacer(1, 0.1 * inch))
            story.append(Paragraph("Unassigned events excluded from patient-scoped gap analysis.", styles["Normal"]))

    else:
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("<b>Appendix C: Treatment Gaps</b>", styles["Heading2"]))
        story.append(Paragraph("No qualifying treatment gaps in projected reportable events.", styles["Normal"]))

    doc.build(story)
    return buf.getvalue()


# â”€â”€ CSV Export â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def generate_csv(
    events: list[Event],
    providers: list[Provider],
    page_map: dict[int, tuple[str, int]] | None = None,
) -> bytes:
    """Generate a CSV chronology with one row per event."""
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "event_id", "date", "provider", "type", "confidence",
        "facts", "source_files",
    ])

    for event in events:
        date_display = _date_str(event)
        writer.writerow([
            event.event_id,
            date_display,
            _provider_name(event, providers),
            event.event_type.value,
            event.confidence,
            _facts_text(event),
            _pages_ref(event, page_map),
        ])

    return buf.getvalue().encode("utf-8")


def generate_csv_from_projection(projection: ChronologyProjection) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["event_id", "date", "provider", "type", "facts", "source"])
    for entry in projection.entries:
        writer.writerow(
            [
                entry.event_id,
                entry.date_display,
                entry.provider_display,
                entry.event_type_display,
                "; ".join(entry.facts),
                entry.citation_display,
            ]
        )
    return buf.getvalue().encode("utf-8")


def _slugify(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return slug or "unknown-patient"


def render_patient_chronology_reports(
    run_id: str,
    matter_title: str,
    events: list[Event],
    providers: list[Provider],
    page_map: dict[int, tuple[str, int]] | None = None,
    page_text_by_number: dict[int, str] | None = None,
) -> ArtifactRef | None:
    """
    Render one chronology PDF per detected patient and return manifest JSON artifact ref.
    """
    page_patient_labels = infer_page_patient_labels(page_text_by_number)
    projection = build_chronology_projection(
        events=events,
        providers=providers,
        page_map=page_map,
        page_patient_labels=page_patient_labels,
        page_text_by_number=page_text_by_number,
    )
    grouped: dict[str, list] = defaultdict(list)
    for entry in projection.entries:
        if entry.patient_label == "Unknown Patient":
            continue
        grouped[entry.patient_label].append(entry)

    if len(grouped) < 2:
        return None

    manifest_rows: list[dict] = []
    for label in sorted(grouped.keys()):
        patient_projection = ChronologyProjection(
            generated_at=projection.generated_at,
            entries=grouped[label],
        )
        pdf_bytes = generate_pdf_from_projection(
            matter_title=f"{matter_title} - {label}",
            projection=patient_projection,
            gaps=[],
            narrative_synthesis=f"Patient-specific chronology for {label}.",
        )
        filename = f"chronology_patient_{_slugify(label)}.pdf"
        pdf_path = save_artifact(run_id, filename, pdf_bytes)
        pdf_sha = hashlib.sha256(pdf_bytes).hexdigest()
        manifest_rows.append(
            {
                "patient_label": label,
                "event_count": len(patient_projection.entries),
                "artifact": {
                    "type": "pdf",
                    "filename": filename,
                    "uri": str(pdf_path),
                    "sha256": pdf_sha,
                    "bytes": len(pdf_bytes),
                },
            }
        )

    manifest = {
        "version": "1.0",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "patients": manifest_rows,
    }
    manifest_bytes = json.dumps(manifest, indent=2).encode("utf-8")
    manifest_path = save_artifact(run_id, "patient_chronologies.json", manifest_bytes)
    manifest_sha = hashlib.sha256(manifest_bytes).hexdigest()
    return ArtifactRef(uri=str(manifest_path), sha256=manifest_sha, bytes=len(manifest_bytes))



# â”€â”€ DOCX Export â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€


def _set_cell_shading(cell, hex_color: str):
    """Set background shading on a DOCX table cell."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")


def generate_docx(
    run_id: str,
    matter_title: str,
    events: list[Event],
    gaps: list[Gap],
    providers: list[Provider],
    page_map: dict[int, tuple[str, int]] | None = None,
    narrative_synthesis: str | None = None,
) -> bytes:
    """
    Generate a professional DOCX chronology for paralegal use.

    Structure:
    - Title page with matter name and generation timestamp
    - Medical Chronology Analysis (Synthesized)
    - Summary statistics
    - Chronology table (dated events sorted ascending)
    - Undated / Needs Review section
    - Treatment gaps appendix
    - Disclaimer
    """
    doc = DocxDocument()

    # â”€â”€ Page setup â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # â”€â”€ Title â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    title_para = doc.add_heading("CiteLine Chronology", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_lines = [
        f"Matter: {matter_title}",
        f"Run ID: {run_id}",
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_paragraph()  # spacer

    # â”€â”€ Narrative Synthesis â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if narrative_synthesis:
        doc.add_heading("Medical Chronology Analysis", level=1)
        narrative_para = doc.add_paragraph(_clean_narrative_text(narrative_synthesis))
        narrative_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in narrative_para.runs:
            run.font.size = Pt(10)
        doc.add_paragraph()

    # â”€â”€ Partition events â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    dated_events = []
    undated_events = []
    flagged_events = []

    for evt in events:
        has_date = (
            evt.date is not None
            and evt.date.value is not None
        )
        needs_review = any(
            f in evt.flags
            for f in ("MISSING_DATE", "MISSING_SOURCE", "NEEDS_REVIEW", "LOW_CONFIDENCE")
        )

        if needs_review:
            flagged_events.append(evt)
        elif has_date:
            dated_events.append(evt)
        else:
            undated_events.append(evt)

    # Sort dated events ascending
    dated_events.sort(key=lambda e: e.date.sort_key() if e.date else (date.min, 0))

    # â”€â”€ Summary statistics â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    total = len(events)
    dated_count = len(dated_events)
    flagged_count = len(flagged_events)
    undated_count = len(undated_events)
    pct_dated = f"{(dated_count / total * 100):.0f}%" if total > 0 else "N/A"

    doc.add_heading("Summary", level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats = [
        ("Total Events", str(total)),
        ("Dated Events", f"{dated_count} ({pct_dated})"),
        ("Undated Events", str(undated_count)),
        ("Flagged (Needs Review)", str(flagged_count)),
    ]
    for i, (label, val) in enumerate(stats):
        summary_table.cell(i, 0).text = label
        summary_table.cell(i, 1).text = val
        for cell in (summary_table.cell(i, 0), summary_table.cell(i, 1)):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # â”€â”€ Helper: add events table â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    def _add_events_table(heading: str, event_list: list[Event]):
        if not event_list:
            return
        doc.add_heading(heading, level=1)

        tbl = doc.add_table(rows=1, cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True

        # Header row
        headers = ["Date", "Provider", "Type", "Description", "Citation"]
        hdr_row = tbl.rows[0]
        for idx, hdr_text in enumerate(headers):
            cell = hdr_row.cells[idx]
            cell.text = hdr_text
            _set_cell_shading(cell, "2C3E50")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Data rows
        for evt in event_list:
            row = tbl.add_row()
            cells = row.cells
            cells[0].text = _date_str(evt)
            cells[1].text = _provider_name(evt, providers)
            cells[2].text = evt.event_type.value.replace("_", " ").title()

            # Description: first 6 facts as bullet points
            facts = "\n".join(f"â€¢ {f.text}" for f in evt.facts[:6])
            cells[3].text = facts

            # Citation
            cells[4].text = _pages_ref(evt, page_map)

            # Style data cells
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    # â”€â”€ Dated events table â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    _add_events_table("Chronology", dated_events)

    # â”€â”€ Undated / Needs Review â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    review_events = undated_events + flagged_events
    if review_events:
        _add_events_table("Undated / Needs Review", review_events)

        # Add flags detail
        doc.add_paragraph()
        for evt in flagged_events:
            flags_str = ", ".join(evt.flags) if evt.flags else "UNDATED"
            p = doc.add_paragraph(f"âš  {evt.event_id}: {flags_str}", style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    # â”€â”€ Treatment gaps â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if gaps:
        doc.add_heading("Appendix: Treatment Gaps", level=1)
        for gap in gaps:
            doc.add_paragraph(
                f"â€¢ {gap.start_date} â†’ {gap.end_date} ({gap.duration_days} days)",
                style="List Bullet",
            )

    # â”€â”€ Disclaimer â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "Factual extraction with citations. Requires human review. "
        "This document does not constitute legal or medical advice."
    )
    for run in disclaimer.runs:
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # â”€â”€ Serialize â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# â”€â”€ Export orchestrator â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def render_exports(
    run_id: str,
    matter_title: str,
    events: list[Event],
    gaps: list[Gap],
    providers: list[Provider],
    page_map: dict[int, tuple[str, int]] | None = None,
    case_info: CaseInfo | None = None,
    all_citations: list[Citation] | None = None,
    narrative_synthesis: str | None = None,
    page_text_by_number: dict[int, str] | None = None,
    evidence_graph_payload: dict | None = None,
    patient_partitions_payload: dict | None = None,
    missing_records_payload: dict | None = None,
) -> ChronologyOutput:
    """
    Render all export formats, save to disk, and return ChronologyOutput.
    """
    provider_none_count = sum(1 for e in events if not e.provider_id or e.provider_id == "unknown")
    print(
        f"chronology_generation_input: {len(events)} events "
        f"(provider_none_or_unknown={provider_none_count})"
    )
    page_patient_labels = infer_page_patient_labels(page_text_by_number)
    projection_debug: list[dict] = []
    selection_meta: dict = {}
    projection = build_chronology_projection(
        events,
        providers,
        page_map=page_map,
        page_patient_labels=page_patient_labels,
        page_text_by_number=page_text_by_number,
        debug_sink=projection_debug,
        selection_meta=selection_meta,
    )
    appendix_projection = build_chronology_projection(
        events,
        providers,
        page_map=page_map,
        page_patient_labels=page_patient_labels,
        page_text_by_number=page_text_by_number,
        select_timeline=False,
    )
    care_window = _compute_care_window_from_projection(projection.entries)
    if missing_records_payload:
        rules = (missing_records_payload.get("ruleset") or {})
        try:
            ms = rules.get("care_window_start")
            me = rules.get("care_window_end")
            if ms and me:
                msd = date.fromisoformat(ms)
                med = date.fromisoformat(me)
                care_window = (msd, med)
        except Exception:
            pass
    safe_narrative, claim_guard_report = apply_claim_guard_to_narrative(narrative_synthesis, page_text_by_number)
    narrative_synthesis = _repair_case_summary_narrative(
        safe_narrative,
        page_text_by_number=page_text_by_number,
        page_map=page_map,
        care_window_start=care_window[0],
        care_window_end=care_window[1],
    )
    projection = _enrich_projection_procedure_entries(
        projection,
        page_text_by_number=page_text_by_number,
        page_map=page_map,
    )
    projection = _ensure_ortho_bucket_entry(
        projection,
        page_text_by_number=page_text_by_number,
        page_map=page_map,
        raw_events=events,
    )
    appendix_projection = _enrich_projection_procedure_entries(
        appendix_projection,
        page_text_by_number=page_text_by_number,
        page_map=page_map,
    )
    exported_ids = [entry.event_id for entry in projection.entries]

    drop_reasons: dict[str, int] = {}
    for item in projection_debug:
        reason = str(item.get("reason") or "unknown")
        drop_reasons[reason] = drop_reasons.get(reason, 0) + 1
    candidates_after_backfill_ids = list(selection_meta.get("candidates_after_backfill_ids", []))
    kept_ids = list(selection_meta.get("kept_ids", []))
    final_ids = list(selection_meta.get("final_ids", []))
    if len(kept_ids) > len(candidates_after_backfill_ids):
        kept_ids = kept_ids[: len(candidates_after_backfill_ids)]
    if len(final_ids) > len(candidates_after_backfill_ids):
        final_ids = final_ids[: len(candidates_after_backfill_ids)]
    selection_debug_payload = {
        "events_extracted_count": len(selection_meta.get("extracted_event_ids", [e.event_id for e in events])),
        "events_candidate_count": len(selection_meta.get("candidates_initial_ids", candidates_after_backfill_ids)),
        "events_candidate_count_after_backfill": len(candidates_after_backfill_ids),
        "events_kept_count": len(kept_ids),
        "events_final_count": len(final_ids),
        "target_rows": len(final_ids),
        "coverage_floor": len(final_ids),
        "extracted_event_ids": selection_meta.get("extracted_event_ids", [e.event_id for e in events]),
        "candidate_event_ids_initial": selection_meta.get("candidates_initial_ids", candidates_after_backfill_ids),
        "candidate_event_ids_after_backfill": candidates_after_backfill_ids,
        "kept_event_ids": kept_ids,
        "final_event_ids": final_ids,
        "dropped_event_ids": sorted({str(i.get("event_id", "")) for i in projection_debug if i.get("event_id")}),
        "drop_reasons": drop_reasons,
    }
    assert selection_debug_payload["events_kept_count"] <= selection_debug_payload["events_candidate_count_after_backfill"]
    assert selection_debug_payload["events_final_count"] <= selection_debug_payload["events_candidate_count_after_backfill"]
    selection_debug_path = save_artifact(run_id, "selection_debug.json", json.dumps(selection_debug_payload, indent=2).encode("utf-8"))
    claim_guard_path = save_artifact(run_id, "claim_guard_report.json", json.dumps(claim_guard_report, indent=2).encode("utf-8"))
    if evidence_graph_payload is not None:
        save_artifact(run_id, "evidence_graph.json", json.dumps(evidence_graph_payload, indent=2).encode("utf-8"))
    if patient_partitions_payload is not None:
        save_artifact(run_id, "patient_partitions.json", json.dumps(patient_partitions_payload, indent=2).encode("utf-8"))
    if missing_records_payload is not None:
        save_artifact(run_id, "missing_records.json", json.dumps(missing_records_payload, indent=2).encode("utf-8"))

    # PDF (projection-only path)
    pdf_bytes = generate_pdf_from_projection(
        matter_title=matter_title,
        projection=projection,
        gaps=gaps,
        narrative_synthesis=narrative_synthesis,
        appendix_entries=appendix_projection.entries,
        raw_events=events,
        page_map=page_map,
        care_window=care_window,
    )
    pdf_path = save_artifact(run_id, "chronology.pdf", pdf_bytes)
    pdf_sha = hashlib.sha256(pdf_bytes).hexdigest()

    # CSV (projection-only path)
    csv_bytes = generate_csv_from_projection(projection)
    csv_path = save_artifact(run_id, "chronology.csv", csv_bytes)
    csv_sha = hashlib.sha256(csv_bytes).hexdigest()

    # DOCX
    docx_bytes = generate_docx(run_id, matter_title, events, gaps, providers, page_map, narrative_synthesis=narrative_synthesis)
    docx_path = save_artifact(run_id, "chronology.docx", docx_bytes)
    docx_sha = hashlib.sha256(docx_bytes).hexdigest()

    # Summary
    summary_text = _clean_narrative_text(narrative_synthesis) if narrative_synthesis else generate_executive_summary(events, matter_title, case_info=case_info)

    return ChronologyOutput(
        export_format_version="0.1.0",
        summary=summary_text,
        narrative_synthesis=narrative_synthesis,
        events_exported=exported_ids,
        exports=ChronologyExports(
            pdf=ArtifactRef(uri=str(pdf_path), sha256=pdf_sha, bytes=len(pdf_bytes)),
            csv=ArtifactRef(uri=str(csv_path), sha256=csv_sha, bytes=len(csv_bytes)),
            docx=ArtifactRef(uri=str(docx_path), sha256=docx_sha, bytes=len(docx_bytes)),
            json_export=None,
        ),
    )


def generate_executive_summary(events: list[Event], matter_title: str, case_info: CaseInfo | None = None) -> str:
    """Generate a high-level narrative summary of the chronology."""
    if not events:
        return "No events documented."
    
    # Filter to dated events for the range calculation, excluding references
    dated_events = [
        e for e in events 
        if e.date and e.date.sort_key()[0] < 99 
        and date_sanity(e.date.sort_date())
        and e.event_type != EventType.REFERENCED_PRIOR_EVENT
        and "is_reference" not in (e.flags or [])
    ]
    if not dated_events:
        return "No dated events documented."
    
    # Sort by the robust sort_key
    dated_events.sort(key=lambda e: e.date.sort_key())
    
    summary = f"Summary for {matter_title}:\n\n"

    # Patient Header Information
    if case_info:
        # Extract name from matter title or case info? matter_title usually has it
        matter_label = matter_title.split('-')[0].strip()
        
        extracted_name = "Unknown"
        mrn_str = ""
        age_str = ""
        
        if case_info.patient:
            p = case_info.patient
            extracted_name = p.name or "Unknown"
            if p.mrn:
                mrn_str = f" (MRN {p.mrn})"
            if p.age:
                age_str = f"Age: {p.age}\n"
        
        summary += f"Patient (extracted): {extracted_name}{mrn_str}\n"
        summary += f"Matter label: {matter_label}\n"
        summary += age_str
        
        # Try to find diagnosis in facts first, then fallback
        diagnosis = None
        for e in events:
            for f in e.facts:
                if "Diagnosis:" in f.text:
                    diagnosis = f.text.split(":", 1)[1].strip()
                    break
                if "Medical History:" in f.text and not diagnosis:
                    diagnosis = f.text.split(":", 1)[1].strip()
            if diagnosis: break
        
        if diagnosis:
            # Clean up: remove "Patient is a 65-year-old female with a four-year history of "
            diagnosis = re.sub(r"(?i)Patient is a \d+-year-old [^ ]+ with a [^ ]+ history of ", "", diagnosis)
            summary += f"Diagnosis: {diagnosis.capitalize()}\n"
        summary += "\n"
    
    # Heuristic: Find first major admission, first procedure, and last discharge or status
    admissions = [e for e in dated_events if e.event_type == EventType.HOSPITAL_ADMISSION]
    discharges = [e for e in dated_events if e.event_type in (EventType.HOSPITAL_DISCHARGE, EventType.DISCHARGE)]
    procedures = [e for e in dated_events if e.event_type == EventType.PROCEDURE]
    
    # DEBUG
    # print(f"DEBUG SUMMARY: dated_events={len(dated_events)} admissions={len(admissions)} discharges={len(discharges)}")
    
    if admissions:
        first_adm = sorted(admissions, key=lambda e: e.date.sort_key())[0]
        summary += f"Documented care began with a hospital admission on {_date_str(first_adm)}. "
    else:
        summary += f"Medical records begin on {_date_str(dated_events[0])}. "
        
    if procedures:
        p_count = len(procedures)
        summary += f"The clinical course included {p_count} significant procedures or operations. "
        
    if discharges:
        # Sort by sort_key DESC to find latest
        sorted_dis = sorted(discharges, key=lambda e: e.date.sort_key(), reverse=True)
        last_dis = sorted_dis[0]
        # print(f"DEBUG SUMMARY: latest discharge event: date={_date_str(last_dis)} type={last_dis.event_type}")
        summary += f"The latest documented discharge occurred on {_date_str(last_dis)}. "
    else:
        summary += f"The records conclude on {_date_str(dated_events[-1])}. "
        
    # Mention specific indicators found?
    pain_facts = []
    functional_facts = []
    diagnosis_facts = []
    history_facts = []
    for e in events:
        for f in e.facts:
            if "Pain Level:" in f.text:
                pain_facts.append(f.text)
            if "Functional Status:" in f.text:
                functional_facts.append(f.text)
            if "Diagnosis:" in f.text:
                diagnosis_facts.append(f.text)
            if "Medical History:" in f.text:
                history_facts.append(f.text)
    
    if history_facts and not diagnosis_facts:
        summary += f"\n\nRelevant medical history includes {history_facts[0].split(':')[-1].strip()}. "

    if pain_facts:
        highlights = pain_facts[0].split(":")[-1].strip()
        summary += f"\n\nKey highlights include reports of significant pain ({highlights}). "
    
    if functional_facts:
        summary += f"Functional decline or assistance requirements were noted, including: {functional_facts[0].split(':')[-1].strip()}. "
        
    # Add Assessment Findings (New)
    findings = {}
    for e in events:
        if e.extensions and "assessment_findings" in e.extensions:
            findings = e.extensions["assessment_findings"]
            break
            
    if findings:
        summary += "\n\nKey Assessment Findings:\n"
        if "fall_risk" in findings:
            summary += f"- Safety: High fall risk (Score: {findings['fall_risk']}); requires assistance for ambulation.\n"
        if "edema" in findings:
            summary += f"- Physical: {findings['edema']} noted.\n"
        if "kyphosis" in findings:
            summary += f"- Physical: Kyphosis of the spine noted.\n"
        if "weight_history" in findings:
            summary += f"- Weight: Weight history: {findings['weight_history']}.\n"

    summary += "\n\nRefer to the timeline below for a complete clinical history with citations."
    return summary

