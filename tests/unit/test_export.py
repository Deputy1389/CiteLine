"""
Unit tests for Step 12 — Export rendering (DOCX focus).
"""
from __future__ import annotations

from datetime import date
import json
from pathlib import Path
from uuid import uuid4

import pytest

from packages.shared.models import (
    Event,
    EventDate,
    EventType,
    Fact,
    FactKind,
    Gap,
    Provider,
    ProviderType,
    DateKind,
    DateSource,
)


# ── Fixtures ──────────────────────────────────────────────────────────────

def _make_event(
    event_id: str = "evt-1",
    event_type: EventType = EventType.OFFICE_VISIT,
    with_date: bool = True,
    flags: list[str] | None = None,
    confidence: int = 85,
) -> Event:
    dt = None
    if with_date:
        dt = EventDate(
            kind=DateKind.SINGLE,
            value=date(2024, 3, 15),
            source=DateSource.TIER1,
        )
    return Event(
        event_id=event_id,
        provider_id="prov-1",
        event_type=event_type,
        date=dt,
        facts=[
            Fact(text="Patient presented with lower back pain", kind=FactKind.CHIEF_COMPLAINT, verbatim=True, citation_id="cit-1"),
            Fact(text="Prescribed ibuprofen 600mg", kind=FactKind.PLAN, verbatim=True, citation_id="cit-2"),
        ],
        confidence=confidence,
        flags=flags or [],
        citation_ids=["cit-1", "cit-2"],
        source_page_numbers=[1, 2],
    )


def _make_provider() -> Provider:
    return Provider(
        provider_id="prov-1",
        detected_name_raw="Dr. Smith",
        normalized_name="Dr. Smith, MD",
        provider_type=ProviderType.SPECIALIST,
        confidence=90,
    )


def _make_gap() -> Gap:
    return Gap(
        gap_id="gap-1",
        start_date=date(2024, 3, 15),
        end_date=date(2024, 5, 1),
        duration_days=47,
        threshold_days=45,
        confidence=80,
    )


# ── DOCX Tests ────────────────────────────────────────────────────────────


class TestGenerateDocx:
    def test_produces_valid_docx(self):
        from apps.worker.steps.step12_export import generate_docx

        events = [_make_event()]
        providers = [_make_provider()]
        gaps = [_make_gap()]

        docx_bytes = generate_docx("run-123", "Smith v Jones", events, gaps, providers)
        # DOCX files are ZIP archives starting with PK
        assert len(docx_bytes) > 100
        assert docx_bytes[:2] == b"PK"

    def test_empty_events(self):
        from apps.worker.steps.step12_export import generate_docx

        docx_bytes = generate_docx("run-456", "Empty Case", [], [], [])
        assert len(docx_bytes) > 100
        assert docx_bytes[:2] == b"PK"

    def test_partitions_flagged_events(self):
        from apps.worker.steps.step12_export import generate_docx
        from docx import Document as DocxDocument

        dated = _make_event(event_id="dated-1")
        flagged = _make_event(
            event_id="flagged-1",
            flags=["MISSING_SOURCE", "NEEDS_REVIEW"],
        )
        undated = _make_event(event_id="undated-1", with_date=False)

        docx_bytes = generate_docx(
            "run-789", "Test Case",
            [dated, flagged, undated], [], [_make_provider()],
        )

        # Parse the DOCX and verify sections exist
        import io
        doc = DocxDocument(io.BytesIO(docx_bytes))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Chronology" in headings
        assert "Undated / Needs Review" in headings

    def test_includes_summary_stats(self):
        from apps.worker.steps.step12_export import generate_docx
        from docx import Document as DocxDocument

        events = [_make_event(), _make_event(event_id="evt-2")]
        docx_bytes = generate_docx("run-abc", "Stats Case", events, [], [_make_provider()])

        import io
        doc = DocxDocument(io.BytesIO(docx_bytes))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Summary" in headings

    def test_with_page_map(self):
        from apps.worker.steps.step12_export import generate_docx

        events = [_make_event()]
        page_map = {1: ("medical_records.pdf", 1), 2: ("medical_records.pdf", 2)}

        docx_bytes = generate_docx(
            "run-map", "Provenance Case",
            events, [], [_make_provider()], page_map=page_map,
        )
        assert len(docx_bytes) > 100

    def test_treatment_gaps_section(self):
        from apps.worker.steps.step12_export import generate_docx
        from docx import Document as DocxDocument

        gaps = [_make_gap()]
        docx_bytes = generate_docx("run-gap", "Gap Case", [_make_event()], gaps, [_make_provider()])

        import io
        doc = DocxDocument(io.BytesIO(docx_bytes))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Appendix: Treatment Gaps" in headings

    def test_narrative_mode_still_includes_chronology(self):
        from apps.worker.steps.step12_export import generate_docx
        from docx import Document as DocxDocument

        events = [
            _make_event(event_id="evt-a"),
            _make_event(event_id="evt-b"),
            _make_event(event_id="evt-c"),
        ]
        docx_bytes = generate_docx(
            "run-narrative",
            "Narrative Case",
            events,
            [],
            [_make_provider()],
            narrative_synthesis="Narrative summary",
        )

        import io
        doc = DocxDocument(io.BytesIO(docx_bytes))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Chronology" in headings


# ── PDF sanity ────────────────────────────────────────────────────────────


class TestGeneratePdf:
    def test_produces_valid_pdf(self):
        from apps.worker.steps.step12_export import generate_pdf

        events = [_make_event()]
        providers = [_make_provider()]

        pdf_bytes = generate_pdf("run-pdf", "PDF Case", events, [], providers)
        assert len(pdf_bytes) > 100
        assert pdf_bytes[:5] == b"%PDF-"

    def test_projection_pdf_includes_paralegal_sections(self):
        import fitz
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        from datetime import datetime, timezone

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="evt-1",
                    date_display="2024-03-15 (time not documented)",
                    provider_display="Interim LSU Public Hospital",
                    event_type_display="Imaging Study",
                    patient_label="Alice111 Smith222",
                    facts=["Impression: right shoulder fracture", "Prescribed ibuprofen 600mg"],
                    citation_display="sample.pdf p. 1",
                    confidence=90,
                )
            ],
        )
        pdf_bytes = generate_pdf_from_projection(
            matter_title="Projection Case",
            projection=projection,
            gaps=[],
            narrative_synthesis="Clean narrative",
        )
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        assert "Facility/Clinician:" in text
        assert "What Happened:" in text
        assert "Why It Matters:" in text
        assert "Citation(s):" in text
        assert "Appendix A: Medications" in text
        assert "Appendix B: Diagnoses/Problems" in text
        assert "Appendix C: Treatment Gaps" in text
        assert "Appendix D: Patient-Reported Outcomes" in text

    def test_projection_pdf_includes_med_change_and_disposition_when_present(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="evt-2",
                    date_display="2024-01-10 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Hospital Admission",
                    patient_label="John Doe",
                    facts=[
                        "Medication started: lisinopril 10 mg tablet daily.",
                        "Patient admitted for inpatient care due to dizziness.",
                        "Disposition: discharged home.",
                    ],
                    citation_display="record.pdf p. 2",
                    confidence=88,
                )
            ],
        )
        pdf_bytes = generate_pdf_from_projection(
            matter_title="MedChange Case",
            projection=projection,
            gaps=[],
            narrative_synthesis="Deterministic narrative.",
        )
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        assert "Medication Changes:" in text
        assert "Disposition:" in text

    def test_projection_pdf_detects_patient_reported_outcomes(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="evt-pro",
                    date_display="2024-01-10 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="John Doe",
                    facts=[
                        "How much did pain interfere with your day-to-day activities? Quite a bit.",
                        "PHQ-9: 16",
                    ],
                    citation_display="record.pdf p. 3",
                    confidence=88,
                )
            ],
        )
        pdf_bytes = generate_pdf_from_projection(
            matter_title="PRO Case",
            projection=projection,
            gaps=[],
            narrative_synthesis="Deterministic narrative.",
        )
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        assert "Appendix D: Patient-Reported Outcomes" in text
        assert "pain interfere" in text.lower() or "phq-9" in text.lower()
        assert "No patient-reported outcome measures identified" not in text

    def test_projection_pdf_detects_medication_change_semantics(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="m1",
                    date_display="2023-02-28 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="John Doe",
                    facts=["Hydrocodone/APAP 5 mg tablet prescribed."],
                    citation_display="record.pdf p. 1",
                    confidence=80,
                ),
                ChronologyProjectionEntry(
                    event_id="m2",
                    date_display="2023-06-18 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="John Doe",
                    facts=["12 HR Hydrocodone Bitartrate 10 MG Oral Capsule, Extended Release noted."],
                    citation_display="record.pdf p. 2",
                    confidence=82,
                ),
            ],
        )
        pdf_bytes = generate_pdf_from_projection(
            matter_title="Med Change Case",
            projection=projection,
            gaps=[],
            narrative_synthesis="Deterministic narrative.",
        )
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count)).lower()
        assert "appendix a: medications (material changes)" in text
        assert "hydrocodone" in text
        assert "dose" in text or "formulation changed" in text or "started" in text

    def test_timeline_what_happened_strips_sdoh_noise(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="evt-sdoh",
                    date_display="2024-02-10 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Hospital Discharge",
                    patient_label="Jane Doe",
                    facts=[
                        "Preferred language: English",
                        "Have you been afraid of your partner or ex-partner in the past year?",
                        "Discharged home with follow-up.",
                    ],
                    citation_display="record.pdf p. 9",
                    confidence=90,
                )
            ],
        )
        pdf_bytes = generate_pdf_from_projection("SDOH Guard", projection, gaps=[], narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count)).lower()
        assert "what happened:" in text
        assert "what happened: reason:" in text
        assert "preferred language" not in text.split("appendix f: social determinants/intake")[0]
        assert "afraid of your partner" not in text.split("appendix f: social determinants/intake")[0]

    def test_gap_anchors_use_adjacent_boundary_events(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="a1",
                    date_display="2024-01-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Hospital Admission",
                    patient_label="John Doe",
                    facts=["Hospital admission for acute issue."],
                    citation_display="record.pdf p. 1",
                    confidence=90,
                ),
                ChronologyProjectionEntry(
                    event_id="a2",
                    date_display="2024-08-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="John Doe",
                    facts=["Follow-up visit documented."],
                    citation_display="record.pdf p. 2",
                    confidence=80,
                ),
            ],
        )
        gaps = [
            Gap(
                gap_id="g-adj",
                start_date=date(2024, 1, 1),
                end_date=date(2024, 8, 1),
                duration_days=213,
                threshold_days=180,
                confidence=80,
                related_event_ids=["a1", "a2"],
            )
        ]
        pdf_bytes = generate_pdf_from_projection("Gap Anchor", projection, gaps=gaps, narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        assert "Appendix C1: Gap Boundary Anchors" in text
        assert "Last before gap: 2024-01-01" in text
        assert "First after gap: 2024-08-01" in text

    def test_top10_diversity_rule_surfaces_multiple_buckets(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="t1",
                    date_display="2024-01-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Hospital Admission",
                    patient_label="P1",
                    facts=["Hospital admission documented."],
                    citation_display="r.pdf p. 1",
                    confidence=92,
                ),
                ChronologyProjectionEntry(
                    event_id="t2",
                    date_display="2024-01-02 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Procedure/Surgery",
                    patient_label="P1",
                    facts=["Procedure performed in operating room."],
                    citation_display="r.pdf p. 2",
                    confidence=95,
                ),
                ChronologyProjectionEntry(
                    event_id="t3",
                    date_display="2024-01-03 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Imaging Study",
                    patient_label="P1",
                    facts=["Impression: fracture line persists."],
                    citation_display="r.pdf p. 3",
                    confidence=88,
                ),
            ],
        )
        pdf_bytes = generate_pdf_from_projection("Top10 Diversity", projection, gaps=[], narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count)).lower()
        assert "top 10 case-driving events" in text
        assert "hospital admission" in text
        assert "procedure/surgery" in text
        assert "imaging study" in text

    def test_top10_excludes_routine_gap_and_nonopioid_med_changes(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="r1",
                    date_display="2024-01-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="P4",
                    facts=["Routine follow-up documented.", "Started acetaminophen 325 mg tablet."],
                    citation_display="r.pdf p. 1",
                    confidence=80,
                ),
                ChronologyProjectionEntry(
                    event_id="r2",
                    date_display="2024-08-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="P4",
                    facts=["Routine follow-up documented.", "Stopped acetaminophen 325 mg tablet."],
                    citation_display="r.pdf p. 2",
                    confidence=80,
                ),
            ],
        )
        gaps = [
            Gap(
                gap_id="g-routine",
                start_date=date(2024, 1, 1),
                end_date=date(2024, 8, 1),
                duration_days=213,
                threshold_days=180,
                confidence=80,
                related_event_ids=["r1", "r2"],
            )
        ]
        pdf_bytes = generate_pdf_from_projection("Top10 Filter", projection, gaps=gaps, narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count)).lower()
        top10_slice = text.split("top 10 case-driving events", 1)[1].split("appendix a:", 1)[0]
        assert "routine_continuity_gap" not in top10_slice
        assert "opioid regimen change" not in top10_slice
        assert "acetaminophen" not in top10_slice

    def test_top10_items_all_have_citations(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="c1",
                    date_display="2024-01-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Hospital Admission",
                    patient_label="P5",
                    facts=["Hospital admission documented."],
                    citation_display="r.pdf p. 1",
                    confidence=90,
                ),
                ChronologyProjectionEntry(
                    event_id="c2",
                    date_display="2024-01-02 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Procedure/Surgery",
                    patient_label="P5",
                    facts=["Procedure performed in OR."],
                    citation_display="r.pdf p. 2",
                    confidence=95,
                ),
            ],
        )
        pdf_bytes = generate_pdf_from_projection("Top10 Cites", projection, gaps=[], narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        top10_slice = text.split("Top 10 Case-Driving Events", 1)[1].split("Appendix A:", 1)[0]
        bullet_lines = [ln for ln in top10_slice.splitlines() if ln.strip().startswith("•")]
        assert bullet_lines
        assert top10_slice.count("Citation(s):") >= len(bullet_lines)
        assert "Citation(s): Not available" not in top10_slice

    def test_hospice_rationale_not_cross_patient(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="a-hospice",
                    date_display="2024-01-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Hospital Discharge",
                    patient_label="A",
                    facts=["Admission to hospice documented."],
                    citation_display="a.pdf p. 1",
                    confidence=90,
                ),
                ChronologyProjectionEntry(
                    event_id="b-start",
                    date_display="2024-02-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="B",
                    facts=["Routine follow-up."],
                    citation_display="b.pdf p. 1",
                    confidence=80,
                ),
                ChronologyProjectionEntry(
                    event_id="b-end",
                    date_display="2024-09-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="B",
                    facts=["Routine follow-up."],
                    citation_display="b.pdf p. 2",
                    confidence=80,
                ),
            ],
        )
        # Mixed-patient related ids should never become a hospice continuity tag.
        gaps = [
            Gap(
                gap_id="g-cross",
                start_date=date(2024, 2, 1),
                end_date=date(2024, 9, 1),
                duration_days=213,
                threshold_days=180,
                confidence=80,
                related_event_ids=["a-hospice", "b-end"],
            )
        ]
        pdf_bytes = generate_pdf_from_projection("Hospice Guard", projection, gaps=gaps, narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count)).lower()
        assert "hospice_continuity_break" not in text

    def test_top10_sanitizer_removes_double_periods(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="s1",
                    date_display="2024-01-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Hospital Admission",
                    patient_label="P6",
                    facts=["Hospital admission documented.."],
                    citation_display="s.pdf p. 1",
                    confidence=90,
                )
            ],
        )
        pdf_bytes = generate_pdf_from_projection("Sanitize", projection, gaps=[], narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        top10_slice = text.split("Top 10 Case-Driving Events", 1)[1].split("Appendix A:", 1)[0]
        assert ".." not in top10_slice
        assert "  " not in top10_slice

    def test_gap_collapsing_for_repeated_routine_intervals(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(event_id="g0", date_display="2020-01-01 (time not documented)", provider_display="Unknown", event_type_display="Follow-Up Visit", patient_label="P7", facts=["Routine follow-up."], citation_display="r.pdf p. 1", confidence=80),
                ChronologyProjectionEntry(event_id="g1", date_display="2021-01-06 (time not documented)", provider_display="Unknown", event_type_display="Follow-Up Visit", patient_label="P7", facts=["Routine follow-up."], citation_display="r.pdf p. 2", confidence=80),
                ChronologyProjectionEntry(event_id="g2", date_display="2022-01-11 (time not documented)", provider_display="Unknown", event_type_display="Follow-Up Visit", patient_label="P7", facts=["Routine follow-up."], citation_display="r.pdf p. 3", confidence=80),
                ChronologyProjectionEntry(event_id="g3", date_display="2023-01-16 (time not documented)", provider_display="Unknown", event_type_display="Follow-Up Visit", patient_label="P7", facts=["Routine follow-up."], citation_display="r.pdf p. 4", confidence=80),
            ],
        )
        gaps = [
            Gap(gap_id="ga", start_date=date(2020, 1, 1), end_date=date(2021, 1, 6), duration_days=371, threshold_days=180, confidence=80, related_event_ids=["g0", "g1"]),
            Gap(gap_id="gb", start_date=date(2021, 1, 6), end_date=date(2022, 1, 11), duration_days=370, threshold_days=180, confidence=80, related_event_ids=["g1", "g2"]),
            Gap(gap_id="gc", start_date=date(2022, 1, 11), end_date=date(2023, 1, 16), duration_days=370, threshold_days=180, confidence=80, related_event_ids=["g2", "g3"]),
        ]
        pdf_bytes = generate_pdf_from_projection("Gap Collapse", projection, gaps=gaps, narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        c1 = text.split("Appendix C1: Gap Boundary Anchors", 1)[1].split("Appendix C: Treatment Gaps", 1)[0]
        assert "Repeated annual continuity gaps collapsed" in c1
        assert c1.count("days) [routine_continuity_gap_collapsed]") == 1

    def test_collapsed_gap_uses_outer_boundary_anchors(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(event_id="h0", date_display="2020-01-01 (time not documented)", provider_display="Unknown", event_type_display="Follow-Up Visit", patient_label="P8", facts=["Routine follow-up."], citation_display="r.pdf p. 1", confidence=80),
                ChronologyProjectionEntry(event_id="h1", date_display="2021-01-06 (time not documented)", provider_display="Unknown", event_type_display="Follow-Up Visit", patient_label="P8", facts=["Routine follow-up."], citation_display="r.pdf p. 2", confidence=80),
                ChronologyProjectionEntry(event_id="h2", date_display="2022-01-11 (time not documented)", provider_display="Unknown", event_type_display="Follow-Up Visit", patient_label="P8", facts=["Routine follow-up."], citation_display="r.pdf p. 3", confidence=80),
                ChronologyProjectionEntry(event_id="h3", date_display="2023-01-16 (time not documented)", provider_display="Unknown", event_type_display="Follow-Up Visit", patient_label="P8", facts=["Routine follow-up."], citation_display="r.pdf p. 4", confidence=80),
            ],
        )
        gaps = [
            Gap(gap_id="ha", start_date=date(2020, 1, 1), end_date=date(2021, 1, 6), duration_days=371, threshold_days=180, confidence=80, related_event_ids=["h0", "h1"]),
            Gap(gap_id="hb", start_date=date(2021, 1, 6), end_date=date(2022, 1, 11), duration_days=370, threshold_days=180, confidence=80, related_event_ids=["h1", "h2"]),
            Gap(gap_id="hc", start_date=date(2022, 1, 11), end_date=date(2023, 1, 16), duration_days=370, threshold_days=180, confidence=80, related_event_ids=["h2", "h3"]),
        ]
        pdf_bytes = generate_pdf_from_projection("Gap Anchor Collapse", projection, gaps=gaps, narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        c1 = text.split("Appendix C1: Gap Boundary Anchors", 1)[1].split("Appendix C: Treatment Gaps", 1)[0]
        assert "Last before gap: 2020-01-01" in c1
        assert "First after gap: 2023-01-16" in c1

    def test_inpatient_progress_variant_does_not_repeat_more_than_twice(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(event_id="ip1", date_display="2024-01-01 (time not documented)", provider_display="Unknown", event_type_display="Inpatient Progress", patient_label="P9", facts=["Inpatient progress note."], citation_display="r.pdf p. 1", confidence=88),
                ChronologyProjectionEntry(event_id="ip2", date_display="2024-01-02 (time not documented)", provider_display="Unknown", event_type_display="Inpatient Progress", patient_label="P9", facts=["Inpatient progress note."], citation_display="r.pdf p. 2", confidence=88),
                ChronologyProjectionEntry(event_id="ip3", date_display="2024-01-03 (time not documented)", provider_display="Unknown", event_type_display="Inpatient Progress", patient_label="P9", facts=["Inpatient progress note."], citation_display="r.pdf p. 3", confidence=88),
                ChronologyProjectionEntry(event_id="ip4", date_display="2024-01-04 (time not documented)", provider_display="Unknown", event_type_display="Inpatient Progress", patient_label="P9", facts=["Inpatient progress note."], citation_display="r.pdf p. 4", confidence=88),
            ],
        )
        pdf_bytes = generate_pdf_from_projection("Inpatient Variants", projection, gaps=[], narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count)).lower()
        # Primary inpatient phrase must not repeat more than twice when alternatives exist.
        assert text.count("inpatient course documented; ongoing monitoring and management") <= 2

    def test_disposition_normalization_outputs_canonical_label(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="disp1",
                    date_display="2024-05-12 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Hospital Discharge",
                    patient_label="P2",
                    facts=["Patient discharged home with instructions."],
                    citation_display="r.pdf p. 4",
                    confidence=85,
                )
            ],
        )
        pdf_bytes = generate_pdf_from_projection("Disposition Norm", projection, gaps=[], narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        assert "Disposition: Home" in text

    def test_med_switch_never_emits_incoherent_dual_target(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="m-prev",
                    date_display="2024-01-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="P3",
                    facts=["Medication review: oxycodone 5 mg tablet continued."],
                    citation_display="r.pdf p. 5",
                    confidence=80,
                ),
                ChronologyProjectionEntry(
                    event_id="m-cur",
                    date_display="2024-02-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="P3",
                    facts=["Switched medications: hydrocodone 10 mg ER capsule and oxycodone 5 mg tablet listed."],
                    citation_display="r.pdf p. 6",
                    confidence=82,
                ),
            ],
        )
        pdf_bytes = generate_pdf_from_projection("Med Switch Guard", projection, gaps=[], narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count)).lower()
        assert "appendix a: medications (material changes)" in text
        assert "opioid switch detected (oxycodone -> hydrocodone, oxycodone)" not in text

    def test_no_unknownwhat_seam(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="u1",
                    date_display="2024-01-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Hospital Admission",
                    patient_label="P10",
                    facts=["Hospital admission documented."],
                    citation_display="r.pdf p. 1",
                    confidence=90,
                )
            ],
        )
        pdf_bytes = generate_pdf_from_projection("Seam Check", projection, gaps=[], narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        assert "UnknownWhat Happened" not in text

    def test_no_dot_pdf_spacing_in_pages_ref(self):
        from apps.worker.steps.step12_export import _pages_ref

        evt = _make_event(event_id="p-ref")
        evt.source_page_numbers = [1]
        ref = _pages_ref(evt, page_map={1: ("PAGES . pdf\n", 7)})
        assert ". pdf" not in ref
        assert "PAGES.pdf p. 7" in ref

    def test_top10_dedupe_and_bucket_caps(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        entries = []
        for idx in range(1, 7):
            entries.append(
                ChronologyProjectionEntry(
                    event_id=f"adm{idx}",
                    date_display=f"2024-01-{idx:02d} (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Hospital Admission",
                    patient_label="PX",
                    facts=["Hospital admission documented."],
                    citation_display=f"a.pdf p. {idx}",
                    confidence=80 + idx,
                )
            )
        entries.append(
            ChronologyProjectionEntry(
                event_id="ed1",
                date_display="2024-02-01 (time not documented)",
                provider_display="Unknown",
                event_type_display="Emergency Visit",
                patient_label="PX",
                facts=["Emergency visit documented."],
                citation_display="e.pdf p. 1",
                confidence=90,
            )
        )
        entries.append(
            ChronologyProjectionEntry(
                event_id="proc1",
                date_display="2024-03-01 (time not documented)",
                provider_display="Unknown",
                event_type_display="Procedure/Surgery",
                patient_label="PX",
                facts=["Procedure documented."],
                citation_display="p.pdf p. 1",
                confidence=95,
            )
        )
        projection = ChronologyProjection(generated_at=datetime.now(timezone.utc), entries=entries)
        pdf_bytes = generate_pdf_from_projection("Top10 Caps", projection, gaps=[], narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        top10_slice = text.split("Top 10 Case-Driving Events", 1)[1].split("Appendix A:", 1)[0]
        assert top10_slice.count("Hospital Admission") <= 3
        assert "Citation(s): Not available" not in top10_slice

    def test_med_changes_dedupes_strength_vs_formulation_same_day(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="mc1",
                    date_display="2024-01-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="PM",
                    facts=["Medication review: acetaminophen 325 mg tablet."],
                    citation_display="m.pdf p. 1",
                    confidence=80,
                ),
                ChronologyProjectionEntry(
                    event_id="mc2",
                    date_display="2024-02-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="PM",
                    facts=["Changed to acetaminophen 300 mg ER capsule."],
                    citation_display="m.pdf p. 2",
                    confidence=80,
                ),
                ChronologyProjectionEntry(
                    event_id="mc3",
                    date_display="2024-03-15 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="PM",
                    facts=["Routine follow-up."],
                    citation_display="m.pdf p. 3",
                    confidence=75,
                ),
                ChronologyProjectionEntry(
                    event_id="mc4",
                    date_display="2024-04-15 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="PM",
                    facts=["Routine follow-up."],
                    citation_display="m.pdf p. 4",
                    confidence=75,
                ),
            ],
        )
        pdf_bytes = generate_pdf_from_projection("Med Dedupe", projection, gaps=[], narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count)).lower()
        appendix_a = text.split("appendix a: medications", 1)[1].split("appendix b:", 1)[0]
        feb_lines = [ln for ln in appendix_a.splitlines() if "2024-02-01" in ln and "acetaminophen" in ln]
        assert len(feb_lines) == 1

    def test_sanitizer_strips_colon_dot_and_truncated_fragments(self):
        import fitz
        from datetime import datetime, timezone
        from apps.worker.project.models import ChronologyProjection, ChronologyProjectionEntry
        from apps.worker.steps.step12_export import generate_pdf_from_projection

        projection = ChronologyProjection(
            generated_at=datetime.now(timezone.utc),
            entries=[
                ChronologyProjectionEntry(
                    event_id="sd1",
                    date_display="2024-01-01 (time not documented)",
                    provider_display="Unknown",
                    event_type_display="Follow-Up Visit",
                    patient_label="PS",
                    facts=["Preferred language:.", "Worried about l."],
                    citation_display="s.pdf p. 1",
                    confidence=70,
                )
            ],
        )
        pdf_bytes = generate_pdf_from_projection("Sanitizer Edge", projection, gaps=[], narrative_synthesis="Deterministic narrative.")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join((doc[i].get_text("text") or "") for i in range(doc.page_count))
        assert ":." not in text
        assert "Worried about l." not in text


# ── CSV sanity ────────────────────────────────────────────────────────────


class TestGenerateCsv:
    def test_produces_csv_with_header(self):
        from apps.worker.steps.step12_export import generate_csv

        events = [_make_event()]
        providers = [_make_provider()]

        csv_bytes = generate_csv(events, providers)
        csv_text = csv_bytes.decode("utf-8")

        assert "event_id" in csv_text
        assert "evt-1" in csv_text

    def test_includes_events_with_missing_provider_id(self):
        from apps.worker.steps.step12_export import generate_csv

        event = _make_event()
        event.provider_id = None
        csv_bytes = generate_csv([event], [_make_provider()])
        csv_text = csv_bytes.decode("utf-8")

        assert "evt-1" in csv_text
        assert "Unknown" in csv_text


class TestPatientChronologyReports:
    def test_generates_per_patient_manifest_for_multi_patient_projection(self):
        from apps.worker.steps.step12_export import render_patient_chronology_reports

        evt1 = _make_event(event_id="evt-p1", event_type=EventType.IMAGING_STUDY)
        evt1.facts = [
            Fact(text="Impression: right shoulder fracture with retained fragments", kind=FactKind.IMPRESSION, verbatim=True)
        ]
        evt1.source_page_numbers = [1]
        evt2 = _make_event(event_id="evt-p2", event_type=EventType.IMAGING_STUDY)
        evt2.facts = [
            Fact(text="Impression: left humerus fracture, no dislocation", kind=FactKind.IMPRESSION, verbatim=True)
        ]
        evt2.source_page_numbers = [2]

        run_id = f"test-patient-reports-{uuid4().hex[:8]}"
        ref = render_patient_chronology_reports(
            run_id=run_id,
            matter_title="Multi Patient",
            events=[evt1, evt2],
            providers=[_make_provider()],
            page_map={1: ("sample.pdf", 1), 2: ("sample.pdf", 2)},
            page_text_by_number={
                1: "PATIENT: Alice111 Smith222",
                2: "PATIENT: Bob333 Jones444",
            },
        )
        assert ref is not None
        manifest_path = Path(ref.uri)
        assert manifest_path.exists()
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))
        assert len(payload["patients"]) >= 2
        for row in payload["patients"]:
            assert Path(row["artifact"]["uri"]).exists()
